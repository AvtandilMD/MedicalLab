import os
import sys
import io
import json
import threading
import webbrowser
from datetime import datetime

from flask import Flask, render_template, request, send_file, Response, jsonify
from werkzeug.utils import secure_filename

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ================== PATHS ==================
def get_base_path():
    # EXE და python ორივეზე სწორი მდებარეობა
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))

def get_template_folder():
    # PyInstaller onefile-ში templates მოდის sys._MEIPASS-დან
    if getattr(sys, "frozen", False):
        return os.path.join(sys._MEIPASS, "templates")
    return os.path.join(get_base_path(), "templates")

def get_saved_docs_folder():
    folder = os.path.join(get_base_path(), "saved_docs")
    os.makedirs(folder, exist_ok=True)
    return folder

def get_data_folder():
    folder = os.path.join(get_base_path(), "data")
    os.makedirs(folder, exist_ok=True)
    return folder

def get_database_path():
    return os.path.join(get_base_path(), "patients_db.json")

def get_settings_path():
    return os.path.join(get_base_path(), "settings.json")


app = Flask(__name__, template_folder=get_template_folder())
app.config["MAX_CONTENT_LENGTH"] = 2 * 1024 * 1024  # 2MB signature upload limit


# ================== PATIENT DB ==================
def load_database():
    path = get_database_path()
    if os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {"patients": []}

def save_database(db):
    with open(get_database_path(), "w", encoding="utf-8") as f:
        json.dump(db, f, ensure_ascii=False, indent=2)

def add_patient_record(first_name, last_name, age, test_type, filename, test_date):
    db = load_database()
    rec = {
        "id": (db["patients"][-1]["id"] + 1) if db["patients"] else 1,
        "first_name": first_name or "",
        "last_name": last_name or "",
        "age": age or "",
        "test_type": test_type,
        "filename": filename,
        "test_date": test_date or "",
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }
    db["patients"].append(rec)
    save_database(db)
    return rec

def search_patients(q: str):
    q = (q or "").strip().lower()
    if len(q) < 2:
        return []
    db = load_database()
    out = []
    for p in db["patients"]:
        fn = (p.get("first_name") or "").lower()
        ln = (p.get("last_name") or "").lower()
        if q in fn or q in ln or q in f"{fn} {ln}":
            out.append(p)
    out.sort(key=lambda x: x.get("created_at", ""), reverse=True)
    return out


# ================== SIGNATURE (PERSISTENT) ==================
ALLOWED_SIGNATURE_EXT = {".png", ".jpg", ".jpeg", ".webp"}

def load_settings():
    path = get_settings_path()
    if os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {"signature_file": None}

def save_settings(settings: dict):
    with open(get_settings_path(), "w", encoding="utf-8") as f:
        json.dump(settings, f, ensure_ascii=False, indent=2)

def get_signature_file_path():
    settings = load_settings()
    fname = settings.get("signature_file")
    if not fname:
        return None
    fpath = os.path.join(get_data_folder(), fname)
    return fpath if os.path.exists(fpath) else None

def signature_exists():
    return get_signature_file_path() is not None

def signature_img_url():
    # cache bust
    return f"/signature?v={int(datetime.now().timestamp())}"

def signature_html_img(height_px=55):
    if not signature_exists():
        return "_____________________"
    return f'<img src="{signature_img_url()}" style="height:{height_px}px;object-fit:contain;">'

def add_signature_line_to_docx(doc: Document, width_cm: float = 3.2):
    """
    აბსოლუტურად ყველა DOCX-ის ბოლოს დაამატებს:
      "ხელმოწერა:" + სურათი (თუ არსებობს) ან ხაზები
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)

    run_label = p.add_run("ხელმოწერა: ")
    run_label.bold = True
    run_label.font.size = Pt(10)

    sig_path = get_signature_file_path()
    if sig_path:
        run_pic = p.add_run()
        run_pic.add_picture(sig_path, width=Cm(width_cm))
    else:
        run_line = p.add_run("_____________________")
        run_line.font.size = Pt(10)


# ================== DOCX HELPERS ==================
def set_cell_shading(cell, color_hex):
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_text(cell, text, font_pt=10, bold=False):
    cell.text = text
    # cell.text ქმნის 1 paragraph/1 run-ს
    try:
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(font_pt)
        run.bold = bold
        cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    except Exception:
        pass


# ================== TEMPLATES DATA ==================
CBC_TEMPLATE = {
    "cbc_analysis": [
        {"abbr": "WBC", "parameter": "ლეიკოციტი", "reference_range": "მ. 5.0-10.0; ქ. 5.0-10.0", "unit": "10^9/L"},
        {"abbr": "RBC", "parameter": "ერითროციტი", "reference_range": "მ. 4.5-5.5; ქ. 4.5-5.5", "unit": "10^12/L"},
        {"abbr": "HGB", "parameter": "ჰემოგლობინი", "reference_range": "მ. 140-174; ქ. 120-174", "unit": "g/L"},
        {"abbr": "HCT", "parameter": "ჰემატოკრიტი", "reference_range": "მ. 36-52; ქ. 45-52", "unit": "%"},
        {"abbr": "PLT", "parameter": "თრომბოციტი", "reference_range": "მ. 150-400; ქ. 150-400", "unit": "10^9/L"},
        {"abbr": "RET", "parameter": "რეტიკულოციტი", "reference_range": "მ. 2-10; ქ. 2-10", "unit": "%"},
        {"abbr": "MCV", "parameter": "ერითროც. საშუალო მოცულობა", "reference_range": "მ. 84-96; ქ. 76-96", "unit": "FL"},
        {"abbr": "MCH", "parameter": "HGB საშუალო შემცველობა", "reference_range": "მ. 27-32; ქ. 27-32", "unit": "pg"},
        {"abbr": "MCHC", "parameter": "HGB საშუალო კონცენტრაცია", "reference_range": "მ. 300-350; ქ. 300-350", "unit": "g/l"},
        {"abbr": "RDW", "parameter": "ერითროც. განაწილების ფართი", "reference_range": "მ. 20-42; ქ. 20-42", "unit": "%"},
        {"abbr": "MPV", "parameter": "თრომბოც. საშუალო მოცულობა", "reference_range": "მ. 8-15; ქ. 8-15", "unit": "FL"},
        {"abbr": "PDW", "parameter": "თრომბოც. განაწილების ფართი", "reference_range": "მ. - ; ქ. -", "unit": "%"},
        {"abbr": "ESR", "parameter": "ერითროც. დალექვის სიჩქარე", "reference_range": "მ. 2-10; ქ. 2-15", "unit": "მმ/სთ"},
    ],
    "leukocyte_formula": [
        {"parameter": "მიელოციტი (MIEL %)", "norm": "0%"},
        {"parameter": "მეტამიელოციტი (METAM %)", "norm": "0%"},
        {"parameter": "ჩხირბირთვიანი ნეიტროფილი (Rod NEUT %)", "norm": "0-6%"},
        {"parameter": "სეგმენტბირთვიანი ნეიტროფილი (SEG %)", "norm": "47-72%"},
        {"parameter": "ეოზინოფილი (EO %)", "norm": "0.5-5%"},
        {"parameter": "ბაზოფილი (BASO %)", "norm": "0-1%"},
        {"parameter": "ლიმფოციტი (LYMPH %)", "norm": "19-37%"},
        {"parameter": "მონოციტი (MONO %)", "norm": "3-11%"},
        {"parameter": "პლაზმური უჯრედი (PLAZ %)", "norm": "0.5-1%"},
    ],
}

URINE_TEMPLATE = {
    "header": {"subtitle": "საოჯახო მედიცინის ცენტრი", "phones": ["558-27-55-51", "577-03-97-70"]},
    "test_info": {"code": "UR.7", "name": "შარდის საერთო ანალიზი"},
    "physico_chemical": [
        {"abbr": "", "parameter": "რაოდენობა", "norm": "", "unit": "მლ"},
        {"abbr": "", "parameter": "ფერი", "norm": "ჩალისფერი", "unit": ""},
        {"abbr": "", "parameter": "გამჭვირვალობა", "norm": "გამჭვირვალე", "unit": ""},
        {"abbr": "SG", "parameter": "ხვედრითი წონა", "norm": "1.005-1.030", "unit": ""},
        {"abbr": "PH", "parameter": "რეაქცია", "norm": "5.0-8.0", "unit": ""},
        {"abbr": "PRO", "parameter": "ცილა", "norm": "0", "unit": "g/l"},
        {"abbr": "GLU", "parameter": "გლუკოზა", "norm": "0", "unit": "mmol/l"},
        {"abbr": "KET", "parameter": "კეტონები", "norm": "0", "unit": "mmol/l"},
        {"abbr": "UBG", "parameter": "ურობილინოგენი", "norm": "3.4-17.0", "unit": "µmol/l"},
        {"abbr": "BIL", "parameter": "ბილირუბინი", "norm": "0", "unit": "µmol/l"},
        {"abbr": "NIT", "parameter": "ნიტრატები", "norm": "NEG", "unit": ""},
        {"abbr": "LEU", "parameter": "ლეიკოციტები", "norm": "-", "unit": "Leu/µL"},
        {"abbr": "BLD", "parameter": "ერითროციტები", "norm": "-", "unit": "Ery/µL"},
    ],
    "microscopy": {
        "epithelium": [
            {"key": "squamous", "label": "ბრტყელი"},
            {"key": "transitional", "label": "გარდამავალი"},
            {"key": "renal", "label": "თირკმლის"},
        ],
        "cylinders": [
            {"key": "hyaline", "label": "ჰიალინური"},
            {"key": "granular", "label": "მარცვლოვანი"},
            {"key": "waxy", "label": "ცვილისებური"},
        ],
        "others": [
            {"key": "mucus", "parameter": "ლორწო"},
            {"key": "salts", "parameter": "მარილები"},
            {"key": "bacteria", "parameter": "ბაქტერიები"},
            {"key": "fungi", "parameter": "სოკო"},
        ],
    },
    "footer": {"equipment": "SIEMENS CLINITEK Status+"},
}

CRP_TEMPLATE = {
    "clinic_info": {"description": "საოჯახო მედიცინის ცენტრი", "phones": ["558-27-55-51", "577-03-97-70"]},
    "test_details": {"title_ge": "მაღალი მგრძნობელობის C-რეაქტიული ცილა (BL.7.9.1)"},
    "test_results": [
        {"code": "CRP", "parameter": "C-რეაქტიული ცილა", "reference_range": "0-10", "unit": "mg/L (მგ/ლ)"},
        {"code": "hsCRP", "parameter": "მაღალი მგრძნობელობის C-რეაქტიული ცილა", "reference_range": "0-1", "unit": "mg/L (მგ/ლ)"},
    ],
}

TROPONIN_TEMPLATE = {
    "document_info": {"clinic_description": "საოჯახო მედიცინის ცენტრი", "contact": "ტელ: 577-03-97-70"},
    "test_info": {
        "title": "ტროპონინის ტესტი (BL.7.8)",
        "results_table": [{"code": "BL.7.8", "parameter": "ტროპონინი", "reference_range": "უარყოფითი"}],
    },
    "footer_note": {"equipment": "გამოკვლევა ჩატარდა ანალიზატორ Firance FS-113 _ზე"},
}


# ================== DOCX GENERATORS ==================
def create_cbc_document(fd: dict) -> Document:
    """
    CBC DOCX - კომპაქტური 10pt შრიფტი, რომ ერთ გვერდზე დაეტიოს.
    """
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(0.8)
        s.bottom_margin = Cm(0.5)
        s.left_margin = Cm(1.2)
        s.right_margin = Cm(1.2)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("PREMIUM MEDI / პრემიუმ მედი")
    r.font.size = Pt(14)
    r.bold = True
    r.font.color.rgb = RGBColor(0, 100, 0)
    h.paragraph_format.space_after = Pt(0)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("საოჯახო მედიცინის ცენტრი | ტელ: 577-03-97-70")
    sub_run.font.size = Pt(9)
    sub.paragraph_format.space_after = Pt(4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = title.add_run("BL6 - სისხლის საერთო ანალიზი CBC")
    t.font.size = Pt(11)
    t.bold = True
    title.paragraph_format.space_after = Pt(6)

    info = doc.add_paragraph()
    info.paragraph_format.space_after = Pt(4)
    info.add_run("პაციენტი: ").bold = True
    info.add_run(f"{fd.get('first_name','')} {fd.get('last_name','')}, {fd.get('age','')} წ.   ")
    info.add_run("თარიღი: ").bold = True
    info.add_run(fd.get("test_date", ""))
    for rr in info.runs:
        rr.font.size = Pt(10)

    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run("სისხლის საერთო ანალიზი")
    r1.bold = True
    r1.font.size = Pt(10)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["აბრევ.", "პარამეტრი", "შედეგი", "ნორმა", "ერთ."]
    for i, htxt in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, htxt, font_pt=10, bold=True)
        set_cell_shading(cell, "D9E2F3")

    for item in CBC_TEMPLATE["cbc_analysis"]:
        row = table.add_row()
        set_cell_text(row.cells[0], item["abbr"], 10, False)
        set_cell_text(row.cells[1], item["parameter"], 10, False)
        set_cell_text(row.cells[2], fd.get(f"cbc_{item['abbr']}", ""), 10, False)
        set_cell_text(row.cells[3], item["reference_range"], 10, False)
        set_cell_text(row.cells[4], item["unit"], 10, False)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run("ლეიკოციტარული ფორმულა")
    r2.bold = True
    r2.font.size = Pt(10)

    lt = doc.add_table(rows=1, cols=3)
    lt.style = "Table Grid"
    lheaders = ["პარამეტრი", "შედეგი", "ნორმა"]
    for i, htxt in enumerate(lheaders):
        cell = lt.rows[0].cells[i]
        set_cell_text(cell, htxt, font_pt=10, bold=True)
        set_cell_shading(cell, "E2F0D9")

    for idx, item in enumerate(CBC_TEMPLATE["leukocyte_formula"]):
        row = lt.add_row()
        set_cell_text(row.cells[0], item["parameter"], 10, False)
        set_cell_text(row.cells[1], fd.get(f"leuko_{idx}", ""), 10, False)
        set_cell_text(row.cells[2], item["norm"], 10, False)

    morph = doc.add_paragraph()
    morph.paragraph_format.space_before = Pt(4)
    morph.paragraph_format.space_after = Pt(2)
    morph.add_run("ერითროც. მორფოლოგია: ").bold = True
    morph.add_run(fd.get("erythrocyte_morphology", "") + "   ")
    morph.add_run("ლეიკოც. მორფოლოგია: ").bold = True
    morph.add_run(fd.get("leukocyte_morphology", ""))
    for rr in morph.runs:
        rr.font.size = Pt(10)

    footer = doc.add_paragraph()
    footer.paragraph_format.space_before = Pt(6)
    footer.paragraph_format.space_after = Pt(0)
    footer.add_run("გამოკვლევა შეასრულა: ").bold = True
    footer.add_run(fd.get("doctor_name", ""))
    for rr in footer.runs:
        rr.font.size = Pt(10)

    add_signature_line_to_docx(doc, width_cm=3.0)
    return doc


def create_urine_document(fd: dict) -> Document:
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(0.8)
        s.bottom_margin = Cm(0.8)
        s.left_margin = Cm(1.2)
        s.right_margin = Cm(1.2)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("PREMIUM MEDI / პრემიუმ მედი")
    r.font.size = Pt(16)
    r.bold = True
    r.font.color.rgb = RGBColor(0, 100, 0)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"{URINE_TEMPLATE['header']['subtitle']} | ტელ: {', '.join(URINE_TEMPLATE['header']['phones'])}").font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(f"{URINE_TEMPLATE['test_info']['code']} - {URINE_TEMPLATE['test_info']['name']}")
    tr.font.size = Pt(14)
    tr.bold = True

    info = doc.add_paragraph()
    info.add_run("პაციენტი: ").bold = True
    info.add_run(f"{fd.get('first_name','')} {fd.get('last_name','')}, {fd.get('age','')} წ.   ")
    info.add_run("თარიღი: ").bold = True
    info.add_run(fd.get("test_date", ""))
    for rr in info.runs:
        rr.font.size = Pt(12)

    doc.add_paragraph().add_run("ფიზიკო-ქიმიური თვისებები").bold = True

    t1 = doc.add_table(rows=1, cols=5)
    t1.style = "Table Grid"
    headers = ["აბრევ.", "პარამეტრი", "შედეგი", "ნორმა", "ერთ."]
    for i, htxt in enumerate(headers):
        cell = t1.rows[0].cells[i]
        set_cell_text(cell, htxt, 11, True)
        set_cell_shading(cell, "FFF2CC")

    for idx, item in enumerate(URINE_TEMPLATE["physico_chemical"]):
        row = t1.add_row()
        set_cell_text(row.cells[0], item["abbr"], 11, False)
        set_cell_text(row.cells[1], item["parameter"], 11, False)
        set_cell_text(row.cells[2], fd.get(f"phys_{idx}", ""), 11, False)
        set_cell_text(row.cells[3], item["norm"], 11, False)
        set_cell_text(row.cells[4], item["unit"], 11, False)

    doc.add_paragraph().add_run("მიკროსკოპია").bold = True

    mt = doc.add_table(rows=1, cols=4)
    mt.style = "Table Grid"
    mh = ["ეპითელიუმი", "შედეგი", "ცილინდრები", "შედეგი"]
    for i, htxt in enumerate(mh):
        cell = mt.rows[0].cells[i]
        set_cell_text(cell, htxt, 11, True)
        set_cell_shading(cell, "E2EFDA")

    epi = URINE_TEMPLATE["microscopy"]["epithelium"]
    cyl = URINE_TEMPLATE["microscopy"]["cylinders"]
    for i in range(max(len(epi), len(cyl))):
        row = mt.add_row()
        if i < len(epi):
            set_cell_text(row.cells[0], epi[i]["label"], 11, False)
            set_cell_text(row.cells[1], fd.get(f"epi_{epi[i]['key']}", ""), 11, False)
        if i < len(cyl):
            set_cell_text(row.cells[2], cyl[i]["label"], 11, False)
            set_cell_text(row.cells[3], fd.get(f"cyl_{cyl[i]['key']}", ""), 11, False)

    doc.add_paragraph().add_run("სხვა მონაცემები").bold = True

    ot = doc.add_table(rows=1, cols=4)
    ot.style = "Table Grid"
    oh = ["პარამეტრი", "შედეგი", "პარამეტრი", "შედეგი"]
    for i, htxt in enumerate(oh):
        cell = ot.rows[0].cells[i]
        set_cell_text(cell, htxt, 11, True)
        set_cell_shading(cell, "DDEBF7")

    others = URINE_TEMPLATE["microscopy"]["others"]
    for i in range(0, len(others), 2):
        row = ot.add_row()
        set_cell_text(row.cells[0], others[i]["parameter"], 11, False)
        set_cell_text(row.cells[1], fd.get(f"other_{others[i]['key']}", ""), 11, False)
        if i + 1 < len(others):
            set_cell_text(row.cells[2], others[i + 1]["parameter"], 11, False)
            set_cell_text(row.cells[3], fd.get(f"other_{others[i+1]['key']}", ""), 11, False)

    footer = doc.add_paragraph()
    footer.paragraph_format.space_before = Pt(10)
    footer.add_run("აპარატურა: ").bold = True
    footer.add_run(URINE_TEMPLATE["footer"]["equipment"] + "   ")
    footer.add_run("შეასრულა: ").bold = True
    footer.add_run(fd.get("doctor_name", ""))
    for rr in footer.runs:
        rr.font.size = Pt(11)

    add_signature_line_to_docx(doc, width_cm=3.2)
    return doc


def create_crp_document(fd: dict) -> Document:
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.2)
        s.bottom_margin = Cm(1.2)
        s.left_margin = Cm(1.8)
        s.right_margin = Cm(1.8)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("PREMIUM MEDI / პრემიუმ მედი")
    r.font.size = Pt(18)
    r.bold = True
    r.font.color.rgb = RGBColor(0, 100, 0)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"{CRP_TEMPLATE['clinic_info']['description']} | ტელ: {', '.join(CRP_TEMPLATE['clinic_info']['phones'])}").font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(CRP_TEMPLATE["test_details"]["title_ge"])
    tr.font.size = Pt(14)
    tr.bold = True

    info = doc.add_paragraph()
    info.add_run("პაციენტი: ").bold = True
    info.add_run(f"{fd.get('first_name','')} {fd.get('last_name','')}, {fd.get('age','')} წ.   ")
    info.add_run("თარიღი: ").bold = True
    info.add_run(fd.get("test_date", ""))
    for rr in info.runs:
        rr.font.size = Pt(12)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["კოდი", "პარამეტრი", "შედეგი", "ნორმა", "ერთეული"]
    for i, htxt in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, htxt, 12, True)
        set_cell_shading(cell, "E8DAEF")

    for item in CRP_TEMPLATE["test_results"]:
        code = item["code"]
        row = table.add_row()
        set_cell_text(row.cells[0], code, 12, False)
        set_cell_text(row.cells[1], item["parameter"], 12, False)
        set_cell_text(row.cells[2], fd.get(f"res_{code}", ""), 12, False)
        set_cell_text(row.cells[3], item["reference_range"], 12, False)
        set_cell_text(row.cells[4], item["unit"], 12, False)

    footer = doc.add_paragraph()
    footer.paragraph_format.space_before = Pt(10)
    footer.add_run("შეასრულა: ").bold = True
    footer.add_run(fd.get("doctor_name", ""))
    for rr in footer.runs:
        rr.font.size = Pt(12)

    add_signature_line_to_docx(doc, width_cm=3.5)
    return doc


def create_troponin_document(fd: dict) -> Document:
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.2)
        s.bottom_margin = Cm(1.2)
        s.left_margin = Cm(1.8)
        s.right_margin = Cm(1.8)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("PREMIUM MEDI / პრემიუმ მედი")
    r.font.size = Pt(18)
    r.bold = True
    r.font.color.rgb = RGBColor(0, 100, 0)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"{TROPONIN_TEMPLATE['document_info']['clinic_description']} | {TROPONIN_TEMPLATE['document_info']['contact']}").font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(TROPONIN_TEMPLATE["test_info"]["title"])
    tr.font.size = Pt(14)
    tr.bold = True

    info = doc.add_paragraph()
    info.add_run("პაციენტი: ").bold = True
    info.add_run(f"{fd.get('first_name','')} {fd.get('last_name','')}, {fd.get('age','')} წ.   ")
    info.add_run("თარიღი: ").bold = True
    info.add_run(fd.get("test_date", ""))
    for rr in info.runs:
        rr.font.size = Pt(12)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["კოდი", "პარამეტრი", "შედეგი", "ნორმა"]
    for i, htxt in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, htxt, 12, True)
        set_cell_shading(cell, "FDEBD0")

    row = table.add_row()
    set_cell_text(row.cells[0], "BL.7.8", 12, False)
    set_cell_text(row.cells[1], "ტროპონინი", 12, False)
    set_cell_text(row.cells[2], fd.get("result_value", ""), 12, False)
    set_cell_text(row.cells[3], "უარყოფითი", 12, False)

    eq = doc.add_paragraph()
    eq.paragraph_format.space_before = Pt(8)
    eq.add_run("აპარატურა: ").bold = True
    eq.add_run(TROPONIN_TEMPLATE["footer_note"]["equipment"])
    for rr in eq.runs:
        rr.font.size = Pt(12)

    footer = doc.add_paragraph()
    footer.add_run("შეასრულა: ").bold = True
    footer.add_run(fd.get("doctor_name", ""))
    for rr in footer.runs:
        rr.font.size = Pt(12)

    add_signature_line_to_docx(doc, width_cm=3.5)
    return doc


# ================== ROUTES (UI) ==================
@app.route("/")
def index():
    return render_template("index.html")

@app.route("/cbc")
def cbc_form():
    return render_template("form_cbc.html", template=CBC_TEMPLATE)

@app.route("/urine")
def urine_form():
    return render_template("form_urinalysis.html", template=URINE_TEMPLATE)

@app.route("/crp")
def crp_form():
    return render_template("form_crp.html", template=CRP_TEMPLATE)

@app.route("/trop")
def trop_form():
    return render_template("form_troponin.html", template=TROPONIN_TEMPLATE)


# ================== ROUTES (SEARCH / FILES) ==================
@app.route("/search")
def search():
    q = request.args.get("q", "")
    return jsonify({"results": search_patients(q)})

@app.route("/download/<path:filename>")
def download_file(filename):
    fpath = os.path.join(get_saved_docs_folder(), filename)
    if not os.path.exists(fpath):
        return "Not Found", 404
    return send_file(fpath, as_attachment=True, download_name=filename)

@app.route("/delete/<int:record_id>", methods=["POST"])
def delete_record(record_id):
    db = load_database()
    for i, rec in enumerate(db["patients"]):
        if rec["id"] == record_id:
            fpath = os.path.join(get_saved_docs_folder(), rec["filename"])
            if os.path.exists(fpath):
                os.remove(fpath)
            db["patients"].pop(i)
            save_database(db)
            return jsonify({"success": True})
    return jsonify({"success": False})


# ================== ROUTES (SIGNATURE) ==================
@app.route("/signature/status")
def signature_status():
    return jsonify({
        "exists": signature_exists(),
        "url": signature_img_url() if signature_exists() else None
    })

@app.route("/signature")
def signature_file():
    sig_path = get_signature_file_path()
    if not sig_path:
        return "Not Found", 404
    return send_file(sig_path, as_attachment=False)

@app.route("/signature/upload", methods=["POST"])
def signature_upload():
    if "signature" not in request.files:
        return jsonify({"success": False, "message": "ფაილი არ აიტვირთა"}), 400

    f = request.files["signature"]
    if not f.filename:
        return jsonify({"success": False, "message": "ფაილის სახელი ცარიელია"}), 400

    filename = secure_filename(f.filename)
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_SIGNATURE_EXT:
        return jsonify({"success": False, "message": "დაშვებულია მხოლოდ: png/jpg/jpeg/webp"}), 400

    save_name = f"doctor_signature{ext}"
    save_path = os.path.join(get_data_folder(), save_name)
    f.save(save_path)

    settings = load_settings()
    settings["signature_file"] = save_name
    save_settings(settings)

    return jsonify({"success": True, "url": signature_img_url()})


# ================== PRINT ROUTES (SAVE DOCX + PRINT HTML) ==================
def _save_docx_and_register(doc: Document, fd: dict, test_type: str, prefix: str):
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    last_name = (fd.get("last_name", "") or "").strip()
    first_name = (fd.get("first_name", "") or "").strip()
    safe = "_".join([x for x in [last_name, first_name] if x]) or "patient"
    filename = f"{prefix}_{safe}_{ts}.docx"
    fpath = os.path.join(get_saved_docs_folder(), filename)
    doc.save(fpath)

    add_patient_record(
        first_name=fd.get("first_name", ""),
        last_name=fd.get("last_name", ""),
        age=fd.get("age", ""),
        test_type=test_type,
        filename=filename,
        test_date=fd.get("test_date", ""),
    )
    return filename

@app.route("/cbc/print", methods=["POST"])
def cbc_print():
    fd = request.form.to_dict()
    doc = create_cbc_document(fd)
    _save_docx_and_register(doc, fd, "CBC", "CBC")

    # print HTML (კომპაქტური)
    html = f"""<!DOCTYPE html><html><head><meta charset="UTF-8"><title>CBC</title>
<style>
@page{{size:A4;margin:10mm}}
body{{font-family:Arial,sans-serif;padding:10px;font-size:11px}}
h1{{color:green;text-align:center;font-size:14px;margin:4px 0}}
h2{{text-align:center;font-size:12px;margin:4px 0}}
h3{{font-size:10px;margin:8px 0 4px 0}}
p{{margin:3px 0}}
table{{width:100%;border-collapse:collapse;margin:5px 0}}
th,td{{border:1px solid #ddd;padding:4px;text-align:left;font-size:9px}}
th{{background:#D9E2F3}}
.leuko th{{background:#E2F0D9}}
</style></head><body>
<h1>PREMIUM MEDI / პრემიუმ მედი</h1>
<p style="text-align:center;font-size:9px">საოჯახო მედიცინის ცენტრი | ტელ: 577-03-97-70</p>
<h2>BL6 - სისხლის საერთო ანალიზი CBC</h2>
<p><b>პაციენტი:</b> {fd.get('first_name','')} {fd.get('last_name','')}, {fd.get('age','')} წ.
&nbsp;&nbsp; <b>თარიღი:</b> {fd.get('test_date','')}</p>

<h3>სისხლის საერთო ანალიზი</h3>
<table><tr><th>აბრევ.</th><th>პარამეტრი</th><th>შედეგი</th><th>ნორმა</th><th>ერთ.</th></tr>"""
    for item in CBC_TEMPLATE["cbc_analysis"]:
        ab = item["abbr"]
        html += f"<tr><td>{ab}</td><td>{item['parameter']}</td><td><b>{fd.get(f'cbc_{ab}','')}</b></td><td>{item['reference_range']}</td><td>{item['unit']}</td></tr>"
    html += "</table>"

    html += """<h3>ლეიკოციტარული ფორმულა</h3>
<table class="leuko"><tr><th>პარამეტრი</th><th>შედეგი</th><th>ნორმა</th></tr>"""
    for idx, item in enumerate(CBC_TEMPLATE["leukocyte_formula"]):
        html += f"<tr><td>{item['parameter']}</td><td><b>{fd.get(f'leuko_{idx}','')}</b></td><td>{item['norm']}</td></tr>"
    html += f"""</table>

<p><b>ერითროც. მორფოლოგია:</b> {fd.get('erythrocyte_morphology','')}
&nbsp;&nbsp; <b>ლეიკოც. მორფოლოგია:</b> {fd.get('leukocyte_morphology','')}</p>

<p><b>შეასრულა:</b> {fd.get('doctor_name','')}
&nbsp;&nbsp; <b>ხელმოწერა:</b> {signature_html_img(55)}</p>

<script>
window.onload=function(){{setTimeout(function(){{window.print()}},500)}}
</script>
</body></html>"""
    return Response(html, mimetype="text/html")


@app.route("/urine/print", methods=["POST"])
def urine_print():
    fd = request.form.to_dict()
    doc = create_urine_document(fd)
    _save_docx_and_register(doc, fd, "Urine", "Urine")

    ph = ", ".join(URINE_TEMPLATE["header"]["phones"])
    html = f"""<!DOCTYPE html><html><head><meta charset="UTF-8"><title>Urinalysis</title>
<style>
@page{{size:A4;margin:10mm}}
body{{font-family:Arial,sans-serif;padding:10px;font-size:11px}}
h1{{color:green;text-align:center;font-size:14px;margin:4px 0}}
h2{{text-align:center;font-size:12px;margin:4px 0}}
h3{{font-size:10px;margin:8px 0 4px 0}}
p{{margin:3px 0}}
table{{width:100%;border-collapse:collapse;margin:5px 0}}
th,td{{border:1px solid #ddd;padding:4px;text-align:left;font-size:9px}}
th{{background:#FFF2CC}}
.micro th{{background:#E2EFDA}}
.other th{{background:#DDEBF7}}
</style></head><body>

<h1>PREMIUM MEDI / პრემიუმ მედი</h1>
<p style="text-align:center;font-size:9px">{URINE_TEMPLATE['header']['subtitle']} | ტელ: {ph}</p>
<h2>{URINE_TEMPLATE['test_info']['code']} - {URINE_TEMPLATE['test_info']['name']}</h2>
<p><b>პაციენტი:</b> {fd.get('first_name','')} {fd.get('last_name','')}, {fd.get('age','')} წ.
&nbsp;&nbsp; <b>თარიღი:</b> {fd.get('test_date','')}</p>

<h3>ფიზიკო-ქიმიური თვისებები</h3>
<table><tr><th>აბრევ.</th><th>პარამეტრი</th><th>შედეგი</th><th>ნორმა</th><th>ერთ.</th></tr>"""
    for idx, item in enumerate(URINE_TEMPLATE["physico_chemical"]):
        html += f"<tr><td>{item['abbr']}</td><td>{item['parameter']}</td><td><b>{fd.get(f'phys_{idx}','')}</b></td><td>{item['norm']}</td><td>{item['unit']}</td></tr>"
    html += "</table>"

    html += """<h3>მიკროსკოპია</h3>
<table class="micro"><tr><th>ეპითელიუმი</th><th>შედეგი</th><th>ცილინდრები</th><th>შედეგი</th></tr>"""
    epi = URINE_TEMPLATE["microscopy"]["epithelium"]
    cyl = URINE_TEMPLATE["microscopy"]["cylinders"]
    for i in range(max(len(epi), len(cyl))):
        el = epi[i]["label"] if i < len(epi) else ""
        ev = fd.get(f"epi_{epi[i]['key']}", "") if i < len(epi) else ""
        cl = cyl[i]["label"] if i < len(cyl) else ""
        cv = fd.get(f"cyl_{cyl[i]['key']}", "") if i < len(cyl) else ""
        html += f"<tr><td>{el}</td><td><b>{ev}</b></td><td>{cl}</td><td><b>{cv}</b></td></tr>"
    html += "</table>"

    # სხვა მონაცემები (ლორწო/მარილები/ბაქტერიები/სოკო)
    html += """<h3>სხვა მონაცემები</h3>
<table class="other"><tr><th>პარამეტრი</th><th>შედეგი</th><th>პარამეტრი</th><th>შედეგი</th></tr>"""
    others = URINE_TEMPLATE["microscopy"]["others"]
    for i in range(0, len(others), 2):
        p1 = others[i]["parameter"]
        v1 = fd.get(f"other_{others[i]['key']}", "")
        p2, v2 = "", ""
        if i + 1 < len(others):
            p2 = others[i + 1]["parameter"]
            v2 = fd.get(f"other_{others[i+1]['key']}", "")
        html += f"<tr><td>{p1}</td><td><b>{v1}</b></td><td>{p2}</td><td><b>{v2}</b></td></tr>"
    html += "</table>"

    html += f"""
<p><b>აპარატურა:</b> {URINE_TEMPLATE['footer']['equipment']}</p>
<p><b>შეასრულა:</b> {fd.get('doctor_name','')}
&nbsp;&nbsp; <b>ხელმოწერა:</b> {signature_html_img(55)}</p>

<script>
window.onload=function(){{setTimeout(function(){{window.print()}},500)}}
</script>
</body></html>"""
    return Response(html, mimetype="text/html")


@app.route("/crp/print", methods=["POST"])
def crp_print():
    fd = request.form.to_dict()
    doc = create_crp_document(fd)
    _save_docx_and_register(doc, fd, "CRP", "CRP")

    ph = ", ".join(CRP_TEMPLATE["clinic_info"]["phones"])
    html = f"""<!DOCTYPE html><html><head><meta charset="UTF-8"><title>CRP</title>
<style>
@page{{size:A4;margin:15mm}}
body{{font-family:Arial,sans-serif;padding:15px;font-size:12px}}
h1{{color:green;text-align:center;font-size:16px;margin:4px 0}}
h2{{text-align:center;font-size:14px;margin:6px 0;color:#8e44ad}}
p{{margin:6px 0;font-size:12px}}
table{{width:100%;border-collapse:collapse;margin:12px 0}}
th,td{{border:1px solid #ddd;padding:8px;text-align:left;font-size:12px}}
th{{background:#E8DAEF}}
</style></head><body>

<h1>PREMIUM MEDI / პრემიუმ მედი</h1>
<p style="text-align:center">{CRP_TEMPLATE['clinic_info']['description']} | ტელ: {ph}</p>
<h2>{CRP_TEMPLATE['test_details']['title_ge']}</h2>
<p><b>პაციენტი:</b> {fd.get('first_name','')} {fd.get('last_name','')}, {fd.get('age','')} წ.
&nbsp;&nbsp; <b>თარიღი:</b> {fd.get('test_date','')}</p>

<table><tr><th>კოდი</th><th>პარამეტრი</th><th>შედეგი</th><th>ნორმა</th><th>ერთეული</th></tr>"""
    for item in CRP_TEMPLATE["test_results"]:
        code = item["code"]
        res = fd.get(f"res_{code}", "")
        html += f"<tr><td><b>{code}</b></td><td>{item['parameter']}</td><td><b>{res}</b></td><td>{item['reference_range']}</td><td>{item['unit']}</td></tr>"
    html += f"""</table>

<p><b>შეასრულა:</b> {fd.get('doctor_name','')}
&nbsp;&nbsp; <b>ხელმოწერა:</b> {signature_html_img(60)}</p>

<script>
window.onload=function(){{setTimeout(function(){{window.print()}},500)}}
</script>
</body></html>"""
    return Response(html, mimetype="text/html")


@app.route("/trop/print", methods=["POST"])
def trop_print():
    fd = request.form.to_dict()
    doc = create_troponin_document(fd)
    _save_docx_and_register(doc, fd, "Troponin", "Trop")

    html = f"""<!DOCTYPE html><html><head><meta charset="UTF-8"><title>Troponin</title>
<style>
@page{{size:A4;margin:15mm}}
body{{font-family:Arial,sans-serif;padding:15px;font-size:12px}}
h1{{color:green;text-align:center;font-size:16px;margin:4px 0}}
h2{{text-align:center;font-size:14px;margin:6px 0;color:#d35400}}
p{{margin:6px 0;font-size:12px}}
table{{width:100%;border-collapse:collapse;margin:12px 0}}
th,td{{border:1px solid #ddd;padding:8px;text-align:left;font-size:12px}}
th{{background:#FDEBD0}}
</style></head><body>

<h1>PREMIUM MEDI / პრემიუმ მედი</h1>
<p style="text-align:center">{TROPONIN_TEMPLATE['document_info']['clinic_description']} | {TROPONIN_TEMPLATE['document_info']['contact']}</p>
<h2>{TROPONIN_TEMPLATE['test_info']['title']}</h2>
<p><b>პაციენტი:</b> {fd.get('first_name','')} {fd.get('last_name','')}, {fd.get('age','')} წ.
&nbsp;&nbsp; <b>თარიღი:</b> {fd.get('test_date','')}</p>

<table><tr><th>კოდი</th><th>პარამეტრი</th><th>შედეგი</th><th>ნორმა</th></tr>
<tr><td><b>BL.7.8</b></td><td>ტროპონინი</td><td><b>{fd.get('result_value','')}</b></td><td>უარყოფითი</td></tr>
</table>

<p><b>აპარატურა:</b> {TROPONIN_TEMPLATE['footer_note']['equipment']}</p>
<p><b>შეასრულა:</b> {fd.get('doctor_name','')}
&nbsp;&nbsp; <b>ხელმოწერა:</b> {signature_html_img(60)}</p>

<script>
window.onload=function(){{setTimeout(function(){{window.print()}},500)}}
</script>
</body></html>"""
    return Response(html, mimetype="text/html")


# ================== START ==================
if __name__ == "__main__":
    # მხოლოდ ერთხელ გახსნას ბრაუზერი
    threading.Timer(1.2, lambda: webbrowser.open("http://127.0.0.1:5000")).start()
    app.run(host="127.0.0.1", port=5000, debug=False, use_reloader=False)