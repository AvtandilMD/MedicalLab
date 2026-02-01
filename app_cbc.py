from flask import Flask, render_template, request, send_file, Response
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import io
import os
import sys
from datetime import datetime


def get_base_path():
    if getattr(sys, 'frozen', False):
        return sys._MEIPASS
    return os.path.dirname(os.path.abspath(__file__))


base_path = get_base_path()
template_folder = os.path.join(base_path, 'templates')
static_folder = os.path.join(base_path, 'static')
app = Flask(__name__, template_folder=template_folder, static_folder=static_folder)

# ქართული ფონტის რეგისტრაცია
font_path = os.path.join(static_folder, 'fonts', 'DejaVuSans.ttf')
if os.path.exists(font_path):
    pdfmetrics.registerFont(TTFont('Georgian', font_path))
    FONT_NAME = 'Georgian'
else:
    FONT_NAME = 'Helvetica'

# CBC შაბლონი
cbc_template = {
    "cbc_analysis": [
        {"abbr": "WBC", "parameter": "ლეიკოციტი", "reference_range": "მ. 5.0-10.0; ქ. 5.0-10.0", "unit": "10^9/L"},
        {"abbr": "RBC", "parameter": "ერითროციტი", "reference_range": "მ. 4.5-5.5; ქ. 4.5-5.5", "unit": "10^12/L"},
        {"abbr": "HGB", "parameter": "ჰემოგლობინი", "reference_range": "მ. 140-174; ქ. 120-174", "unit": "g/L"},
        {"abbr": "HCT", "parameter": "ჰემატოკრიტი", "reference_range": "მ. 36-52; ქ. 45-52", "unit": "%"},
        {"abbr": "PLT", "parameter": "თრომბოციტი", "reference_range": "მ. 150-400; ქ. 150-400", "unit": "10^9/L"},
        {"abbr": "RET", "parameter": "რეტიკულოციტი", "reference_range": "მ. 2-10; ქ. 2-10", "unit": "%"},
        {"abbr": "MCV", "parameter": "ერითროც. საშუალო მოცულობა", "reference_range": "მ. 84-96; ქ. 76-96",
         "unit": "FL"},
        {"abbr": "MCH", "parameter": "HGB საშუალო შემცველობა", "reference_range": "მ. 27-32; ქ. 27-32", "unit": "pg"},
        {"abbr": "MCHC", "parameter": "HGB საშუალო კონცენტრაცია", "reference_range": "მ. 300-350; ქ. 300-350",
         "unit": "g/l"},
        {"abbr": "RDW", "parameter": "ერითროც. განაწილების ფართი", "reference_range": "მ. 20-42; ქ. 20-42",
         "unit": "%"},
        {"abbr": "MPV", "parameter": "თრომბოც. საშუალო მოცულობა", "reference_range": "მ. 8-15; ქ. 8-15", "unit": "FL"},
        {"abbr": "PDW", "parameter": "თრომბოც. განაწილების ფართი", "reference_range": "მ. - ; ქ. -", "unit": "%"},
        {"abbr": "ESR", "parameter": "ერითროც. დალექვის სიჩქარე", "reference_range": "მ. 2-10; ქ. 2-15",
         "unit": "მმ/სთ"}
    ],
    "leukocyte_formula": [
        {"parameter": "მიელოციტი (MIEL %)", "norm": "0%"},
        {"parameter": "მეტამიელოციტი (METAM %)", "norm": "0%"},
        {"parameter": "ჩხირბირთვიანი ნეიტროფილი (Rod NEUT %)", "norm": "0-6%"},
        {"parameter": "სეგმენტბირთვიანი ნეიტროფილი (SEG %)", "norm": "47-72%"},
        {"parameter": "ეოზინოფილი (EO %)", "norm": "0.5-5%"},
        {"parameter": "ბაზოფილი (BASO %)", "norm": "0-1%"},
        {"parameter": "ლიმფოციტი (LYMPH %)", "norm": "19-37%"},
        {"parameter": "მონოციტი (MONO %)", "norm": "3-11%"},
        {"parameter": "პლაზმური უჯრედი (PLAZ %)", "norm": "0.5-1%"}
    ]
}


def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def create_word_document(form_data):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # ჰედერი
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("PREMIUM MEDI / პრემიუმ მედი")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 100, 0)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("საოჯახო მედიცინის ცენტრი | ტელ: 558-27-55-51").font.size = Pt(9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = title.add_run("BL6 - სისხლის საერთო ანალიზი CBC")
    t.font.size = Pt(12)
    t.font.bold = True

    # პაციენტის ინფო
    p = doc.add_paragraph()
    p.add_run("პაციენტი: ").bold = True
    p.add_run(f"{form_data.get('first_name', '')} {form_data.get('last_name', '')}, {form_data.get('age', '')} წ.")

    d = doc.add_paragraph()
    d.add_run("თარიღი: ").bold = True
    d.add_run(form_data.get('test_date', ''))

    # CBC ცხრილი
    doc.add_paragraph().add_run("სისხლის საერთო ანალიზი").bold = True

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    headers = ['აბრევ.', 'პარამეტრი', 'შედეგი', 'ნორმა', 'ერთეული']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        set_cell_shading(table.rows[0].cells[i], 'D9E2F3')

    for item in cbc_template["cbc_analysis"]:
        row = table.add_row()
        row.cells[0].text = item['abbr']
        row.cells[1].text = item['parameter']
        row.cells[2].text = form_data.get(f'cbc_{item["abbr"]}', '')
        row.cells[3].text = item['reference_range']
        row.cells[4].text = item['unit']

    # ლეიკოციტარული ფორმულა
    doc.add_paragraph().add_run("ლეიკოციტარული ფორმულა").bold = True

    lt = doc.add_table(rows=1, cols=3)
    lt.style = 'Table Grid'
    lheaders = ['პარამეტრი', 'შედეგი', 'ნორმა']
    for i, h in enumerate(lheaders):
        lt.rows[0].cells[i].text = h
        set_cell_shading(lt.rows[0].cells[i], 'E2F0D9')

    for idx, item in enumerate(cbc_template["leukocyte_formula"]):
        row = lt.add_row()
        row.cells[0].text = item['parameter']
        row.cells[1].text = form_data.get(f'leuko_{idx}', '')
        row.cells[2].text = item['norm']

    # მორფოლოგია
    doc.add_paragraph()
    m1 = doc.add_paragraph()
    m1.add_run("ერითროციტის მორფოლოგია: ").bold = True
    m1.add_run(form_data.get('erythrocyte_morphology', ''))

    m2 = doc.add_paragraph()
    m2.add_run("ლეიკოციტის მორფოლოგია: ").bold = True
    m2.add_run(form_data.get('leukocyte_morphology', ''))

    # ფუტერი
    doc.add_paragraph()
    f = doc.add_paragraph()
    f.add_run("გამოკვლევა შეასრულა: ").bold = True
    f.add_run(form_data.get('doctor_name', ''))

    doc.add_paragraph().add_run("ხელმოწერა: _________________________")

    return doc


def create_pdf_document(form_data):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=1.5 * cm,
        rightMargin=1.5 * cm,
        topMargin=1 * cm,
        bottomMargin=1 * cm
    )

    styles = getSampleStyleSheet()

    # ქართული სტილები
    georgian_style = ParagraphStyle(
        'Georgian',
        fontName=FONT_NAME,
        fontSize=10,
        leading=14
    )

    header_style = ParagraphStyle(
        'GeorgianHeader',
        fontName=FONT_NAME,
        fontSize=14,
        alignment=1,
        textColor=colors.darkgreen,
        spaceAfter=6
    )

    title_style = ParagraphStyle(
        'GeorgianTitle',
        fontName=FONT_NAME,
        fontSize=12,
        alignment=1,
        spaceAfter=12
    )

    story = []

    # ჰედერი
    story.append(Paragraph("PREMIUM MEDI / პრემიუმ მედი", header_style))
    story.append(Paragraph("საოჯახო მედიცინის ცენტრი | ტელ: 558-27-55-51",
                           ParagraphStyle('Sub', fontName=FONT_NAME, fontSize=9, alignment=1)))
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph("BL6 - სისხლის საერთო ანალიზი CBC", title_style))
    story.append(Spacer(1, 0.3 * cm))

    # პაციენტის ინფო
    name = f"{form_data.get('first_name', '')} {form_data.get('last_name', '')}, {form_data.get('age', '')} წ."
    story.append(Paragraph(f"<b>პაციენტი:</b> {name}", georgian_style))
    story.append(Paragraph(f"<b>თარიღი:</b> {form_data.get('test_date', '')}", georgian_style))
    story.append(Spacer(1, 0.4 * cm))

    # CBC ცხრილი
    story.append(Paragraph("<b>სისხლის საერთო ანალიზი</b>", georgian_style))
    story.append(Spacer(1, 0.2 * cm))

    cbc_data = [['აბრევ.', 'პარამეტრი', 'შედეგი', 'ნორმა', 'ერთეული']]
    for item in cbc_template["cbc_analysis"]:
        cbc_data.append([
            item['abbr'],
            item['parameter'],
            form_data.get(f'cbc_{item["abbr"]}', ''),
            item['reference_range'],
            item['unit']
        ])

    cbc_table = Table(cbc_data, colWidths=[1.5 * cm, 5 * cm, 2 * cm, 4 * cm, 2 * cm])
    cbc_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#D9E2F3')),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('FONTNAME', (0, 0), (-1, -1), FONT_NAME),
        ('ALIGN', (2, 1), (2, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    story.append(cbc_table)
    story.append(Spacer(1, 0.4 * cm))

    # ლეიკოციტარული ფორმულა
    story.append(Paragraph("<b>ლეიკოციტარული ფორმულა</b>", georgian_style))
    story.append(Spacer(1, 0.2 * cm))

    leu_data = [['პარამეტრი', 'შედეგი', 'ნორმა']]
    for idx, item in enumerate(cbc_template["leukocyte_formula"]):
        leu_data.append([
            item['parameter'],
            form_data.get(f'leuko_{idx}', ''),
            item['norm']
        ])

    leu_table = Table(leu_data, colWidths=[8 * cm, 3 * cm, 3 * cm])
    leu_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E2F0D9')),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('FONTNAME', (0, 0), (-1, -1), FONT_NAME),
        ('ALIGN', (1, 1), (1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    story.append(leu_table)
    story.append(Spacer(1, 0.4 * cm))

    # მორფოლოგია
    story.append(
        Paragraph(f"<b>ერითროციტის მორფოლოგია:</b> {form_data.get('erythrocyte_morphology', '')}", georgian_style))
    story.append(
        Paragraph(f"<b>ლეიკოციტის მორფოლოგია:</b> {form_data.get('leukocyte_morphology', '')}", georgian_style))
    story.append(Spacer(1, 0.4 * cm))

    # ფუტერი
    story.append(Paragraph(f"<b>გამოკვლევა შეასრულა:</b> {form_data.get('doctor_name', '')}", georgian_style))
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph("ხელმოწერა: _________________________", georgian_style))

    doc.build(story)
    buffer.seek(0)
    return buffer


@app.route('/')
def index():
    return render_template('form_cbc.html', template=cbc_template)


@app.route('/generate_cbc_doc', methods=['POST'])
def generate_doc():
    form_data = request.form.to_dict()
    doc = create_word_document(form_data)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    filename = f"CBC_{form_data.get('last_name', '')}_{datetime.now().strftime('%Y%m%d')}.docx"
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/generate_cbc_pdf', methods=['POST'])
def generate_pdf():
    form_data = request.form.to_dict()
    buffer = create_pdf_document(form_data)

    # PDF კონტენტი ავტო-ბეჭდვით
    pdf_content = buffer.getvalue()

    html_wrapper = f'''
    <!DOCTYPE html>
    <html>
    <head>
        <title>CBC Report - Print</title>
        <style>
            body, html {{ margin: 0; padding: 0; height: 100%; overflow: hidden; }}
            iframe {{ width: 100%; height: 100%; border: none; }}
        </style>
    </head>
    <body>
        <iframe id="pdfFrame" src="data:application/pdf;base64,{__import__('base64').b64encode(pdf_content).decode()}"></iframe>
        <script>
            window.onload = function() {{
                setTimeout(function() {{
                    var frame = document.getElementById('pdfFrame');
                    frame.contentWindow.focus();
                    frame.contentWindow.print();
                }}, 1000);
            }};
        </script>
    </body>
    </html>
    '''
    return Response(html_wrapper, mimetype='text/html')


# ალტერნატიული PDF მარშრუტი (მხოლოდ PDF ნახვისთვის, ბეჭდვის გარეშე)
@app.route('/view_cbc_pdf', methods=['POST'])
def view_pdf():
    form_data = request.form.to_dict()
    buffer = create_pdf_document(form_data)
    return send_file(
        buffer,
        as_attachment=False,
        download_name="cbc_report.pdf",
        mimetype='application/pdf'
    )


if __name__ == '__main__':
    print("=" * 50)
    print("🩸 CBC აპლიკაცია გაშვებულია")
    print("=" * 50)
    print("🌐 გახსენით: http://127.0.0.1:5000")
    print("=" * 50)
    app.run(debug=False, host='127.0.0.1', port=5000)