from flask import Flask, render_template, request, send_file, Response
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import io
import os
import sys
import base64
from datetime import datetime


def get_base_path():
    if getattr(sys, 'frozen', False):
        return sys._MEIPASS
    return os.path.dirname(os.path.abspath(__file__))


base_path = get_base_path()
template_folder = os.path.join(base_path, 'templates')
static_folder = os.path.join(base_path, 'static')
app = Flask(__name__, template_folder=template_folder, static_folder=static_folder)

# ქართული ფონტის რეგისტრაცია
font_path = os.path.join(static_folder, 'fonts', 'DejaVuSans.ttf')
if os.path.exists(font_path):
    pdfmetrics.registerFont(TTFont('Georgian', font_path))
    FONT_NAME = 'Georgian'
else:
    FONT_NAME = 'Helvetica'

# შარდის შაბლონი
urinalysis_template = {
    "header": {
        "clinic_name": "პრემიუმ მედი",
        "clinic_name_en": "PREMIUM MEDI",
        "subtitle": "საოჯახო მედიცინის ცენტრი",
        "address": "ქ. ხაშური, იმერეთის ქუჩა №2",
        "phones": ["558-27-55-51", "577-03-97-70"]
    },
    "test_info": {
        "code": "UR.7",
        "name": "შარდის საერთო ანალიზი"
    },
    "physico_chemical": [
        {"abbr": "", "parameter": "რაოდენობა", "norm": "", "unit": "მლ"},
        {"abbr": "", "parameter": "ფერი", "norm": "ჩალისფერი", "unit": ""},
        {"abbr": "", "parameter": "გამჭვირვალობა", "norm": "გამჭვირვალე", "unit": ""},
        {"abbr": "SG", "parameter": "ხვედრითი წონა", "norm": "1.005-1.030", "unit": ""},
        {"abbr": "PH", "parameter": "რეაქცია", "norm": "5.0-8.0", "unit": ""},
        {"abbr": "PRO", "parameter": "ცილა", "norm": "0", "unit": "g/l"},
        {"abbr": "GLU", "parameter": "გლუკოზა", "norm": "0", "unit": "mmol/l"},
        {"abbr": "KET", "parameter": "კეტონები", "norm": "0", "unit": "mmol/l"},
        {"abbr": "UBG", "parameter": "ურობილინოგენი", "norm": "3.4-17.0", "unit": "µmol/l"},
        {"abbr": "BIL", "parameter": "ბილირუბინი", "norm": "0", "unit": "µmol/l"},
        {"abbr": "NIT", "parameter": "ნიტრატები", "norm": "NEG", "unit": ""},
        {"abbr": "LEU", "parameter": "ლეიკოციტები", "norm": "-", "unit": "Leu/µL"},
        {"abbr": "BLD", "parameter": "ერითროციტები", "norm": "-", "unit": "Ery/µL"}
    ],
    "microscopy": {
        "epithelium": [
            {"key": "squamous", "label": "ბრტყელი"},
            {"key": "transitional", "label": "გარდამავალი"},
            {"key": "renal", "label": "თირკმლის"}
        ],
        "cylinders": [
            {"key": "hyaline", "label": "ჰიალინური"},
            {"key": "granular", "label": "მარცვლოვანი"},
            {"key": "waxy", "label": "ცვილისებური"}
        ],
        "others": [
            {"key": "mucus", "parameter": "ლორწო"},
            {"key": "salts", "parameter": "მარილები"},
            {"key": "bacteria", "parameter": "ბაქტერიები"},
            {"key": "fungi", "parameter": "სოკო"}
        ]
    },
    "footer": {
        "equipment": "SIEMENS CLINITEK Status+"
    }
}


def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def create_urinalysis_document(form_data):
    doc = Document()

    for s in doc.sections:
        s.top_margin = Cm(0.8)
        s.bottom_margin = Cm(0.8)
        s.left_margin = Cm(1.5)
        s.right_margin = Cm(1.5)

    # ჰედერი
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = header.add_run("PREMIUM MEDI")
    r1.font.size = Pt(14)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0, 100, 0)
    header.add_run(" / პრემიუმ მედი").font.size = Pt(12)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        f"{urinalysis_template['header']['subtitle']} | "
        f"{urinalysis_template['header']['address']} | "
        f"ტელ: {', '.join(urinalysis_template['header']['phones'])}"
    ).font.size = Pt(9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = title.add_run(f"{urinalysis_template['test_info']['code']} - {urinalysis_template['test_info']['name']}")
    t.font.size = Pt(12)
    t.font.bold = True

    # პაციენტის ინფო
    p = doc.add_paragraph()
    p.add_run("პაციენტი: ").bold = True
    p.add_run(f"{form_data.get('first_name', '')} {form_data.get('last_name', '')}, {form_data.get('age', '')} წ.")

    d = doc.add_paragraph()
    d.add_run("თარიღი: ").bold = True
    d.add_run(form_data.get('test_date', ''))

    # ფიზიკო-ქიმიური ცხრილი
    doc.add_paragraph().add_run("ფიზიკო-ქიმიური თვისებები").bold = True

    t1 = doc.add_table(rows=1, cols=5)
    t1.style = 'Table Grid'
    h1 = ['აბრევ.', 'პარამეტრი', 'შედეგი', 'ნორმა', 'ერთეული']
    for i, h in enumerate(h1):
        t1.rows[0].cells[i].text = h
        set_cell_shading(t1.rows[0].cells[i], 'FFF2CC')

    for idx, item in enumerate(urinalysis_template["physico_chemical"]):
        row = t1.add_row()
        row.cells[0].text = item['abbr']
        row.cells[1].text = item['parameter']
        row.cells[2].text = form_data.get(f'phys_{idx}', '')
        row.cells[3].text = item['norm']
        row.cells[4].text = item['unit']

    # მიკროსკოპია
    doc.add_paragraph().add_run("მიკროსკოპია").bold = True

    mt = doc.add_table(rows=1, cols=4)
    mt.style = 'Table Grid'
    mh = mt.rows[0].cells
    mh[0].text = "ეპითელიუმი"
    mh[1].text = "შედეგი"
    mh[2].text = "ცილინდრები"
    mh[3].text = "შედეგი"
    for c in mh:
        set_cell_shading(c, 'E2EFDA')

    epi = urinalysis_template["microscopy"]["epithelium"]
    cyl = urinalysis_template["microscopy"]["cylinders"]
    for i in range(max(len(epi), len(cyl))):
        row = mt.add_row()
        if i < len(epi):
            row.cells[0].text = epi[i]['label']
            row.cells[1].text = form_data.get(f'epi_{epi[i]["key"]}', '')
        if i < len(cyl):
            row.cells[2].text = cyl[i]['label']
            row.cells[3].text = form_data.get(f'cyl_{cyl[i]["key"]}', '')

    # სხვა მონაცემები
    doc.add_paragraph().add_run("სხვა მონაცემები").bold = True

    ot = doc.add_table(rows=1, cols=4)
    ot.style = 'Table Grid'
    oh = ot.rows[0].cells
    oh[0].text = "პარამეტრი"
    oh[1].text = "შედეგი"
    oh[2].text = "პარამეტრი"
    oh[3].text = "შედეგი"
    for c in oh:
        set_cell_shading(c, 'DDEBF7')

    others = urinalysis_template["microscopy"]["others"]
    for i in range(0, len(others), 2):
        row = ot.add_row()
        row.cells[0].text = others[i]['parameter']
        row.cells[1].text = form_data.get(f'other_{others[i]["key"]}', '')
        if i + 1 < len(others):
            row.cells[2].text = others[i + 1]['parameter']
            row.cells[3].text = form_data.get(f'other_{others[i + 1]["key"]}', '')

    # ფუტერი
    doc.add_paragraph()
    eq = doc.add_paragraph()
    eq.add_run("აპარატურა: ").bold = True
    eq.add_run(urinalysis_template["footer"]["equipment"])

    f = doc.add_paragraph()
    f.add_run("გამოკვლევა შეასრულა: ").bold = True
    f.add_run(form_data.get('doctor_name', ''))

    doc.add_paragraph().add_run("ხელმოწერა: _________________________")

    return doc


def create_pdf_document(form_data):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=1.5 * cm,
        rightMargin=1.5 * cm,
        topMargin=1 * cm,
        bottomMargin=1 * cm
    )

    # სტილები
    georgian_style = ParagraphStyle(
        'Georgian',
        fontName=FONT_NAME,
        fontSize=10,
        leading=14
    )

    header_style = ParagraphStyle(
        'GeorgianHeader',
        fontName=FONT_NAME,
        fontSize=14,
        alignment=1,
        textColor=colors.darkgreen,
        spaceAfter=6
    )

    title_style = ParagraphStyle(
        'GeorgianTitle',
        fontName=FONT_NAME,
        fontSize=12,
        alignment=1,
        spaceAfter=12
    )

    story = []

    # ჰედერი
    story.append(Paragraph("PREMIUM MEDI / პრემიუმ მედი", header_style))
    story.append(Paragraph(
        f"{urinalysis_template['header']['subtitle']} | ტელ: {', '.join(urinalysis_template['header']['phones'])}",
        ParagraphStyle('Sub', fontName=FONT_NAME, fontSize=9, alignment=1)
    ))
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph(
        f"{urinalysis_template['test_info']['code']} - {urinalysis_template['test_info']['name']}",
        title_style
    ))
    story.append(Spacer(1, 0.3 * cm))

    # პაციენტის ინფო
    name = f"{form_data.get('first_name', '')} {form_data.get('last_name', '')}, {form_data.get('age', '')} წ."
    story.append(Paragraph(f"<b>პაციენტი:</b> {name}", georgian_style))
    story.append(Paragraph(f"<b>თარიღი:</b> {form_data.get('test_date', '')}", georgian_style))
    story.append(Spacer(1, 0.4 * cm))

    # ფიზიკო-ქიმიური ცხრილი
    story.append(Paragraph("<b>ფიზიკო-ქიმიური თვისებები</b>", georgian_style))
    story.append(Spacer(1, 0.2 * cm))

    phys_data = [['აბრევ.', 'პარამეტრი', 'შედეგი', 'ნორმა', 'ერთეული']]
    for idx, item in enumerate(urinalysis_template["physico_chemical"]):
        phys_data.append([
            item['abbr'],
            item['parameter'],
            form_data.get(f'phys_{idx}', ''),
            item['norm'],
            item['unit']
        ])

    phys_table = Table(phys_data, colWidths=[1.5 * cm, 4 * cm, 2.5 * cm, 3 * cm, 2 * cm])
    phys_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#FFF2CC')),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('FONTNAME', (0, 0), (-1, -1), FONT_NAME),
        ('ALIGN', (2, 1), (2, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
    ]))
    story.append(phys_table)
    story.append(Spacer(1, 0.4 * cm))

    # მიკროსკოპია
    story.append(Paragraph("<b>მიკროსკოპია</b>", georgian_style))
    story.append(Spacer(1, 0.2 * cm))

    # ეპითელიუმი & ცილინდრები
    epi = urinalysis_template["microscopy"]["epithelium"]
    cyl = urinalysis_template["microscopy"]["cylinders"]

    micro_data = [['ეპითელიუმი', 'შედეგი', 'ცილინდრები', 'შედეგი']]
    for i in range(max(len(epi), len(cyl))):
        row = ['', '', '', '']
        if i < len(epi):
            row[0] = epi[i]['label']
            row[1] = form_data.get(f'epi_{epi[i]["key"]}', '')
        if i < len(cyl):
            row[2] = cyl[i]['label']
            row[3] = form_data.get(f'cyl_{cyl[i]["key"]}', '')
        micro_data.append(row)

    micro_table = Table(micro_data, colWidths=[4 * cm, 2.5 * cm, 4 * cm, 2.5 * cm])
    micro_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E2EFDA')),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('FONTNAME', (0, 0), (-1, -1), FONT_NAME),
        ('ALIGN', (1, 1), (1, -1), 'CENTER'),
        ('ALIGN', (3, 1), (3, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    story.append(micro_table)
    story.append(Spacer(1, 0.3 * cm))

    # სხვა მონაცემები
    story.append(Paragraph("<b>სხვა მონაცემები</b>", georgian_style))
    story.append(Spacer(1, 0.2 * cm))

    others = urinalysis_template["microscopy"]["others"]
    others_data = [['პარამეტრი', 'შედეგი', 'პარამეტრი', 'შედეგი']]
    for i in range(0, len(others), 2):
        row = [others[i]['parameter'], form_data.get(f'other_{others[i]["key"]}', ''), '', '']
        if i + 1 < len(others):
            row[2] = others[i + 1]['parameter']
            row[3] = form_data.get(f'other_{others[i + 1]["key"]}', '')
        others_data.append(row)

    others_table = Table(others_data, colWidths=[4 * cm, 2.5 * cm, 4 * cm, 2.5 * cm])
    others_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#DDEBF7')),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('FONTNAME', (0, 0), (-1, -1), FONT_NAME),
        ('ALIGN', (1, 1), (1, -1), 'CENTER'),
        ('ALIGN', (3, 1), (3, -1), 'CENTER'),
    ]))
    story.append(others_table)
    story.append(Spacer(1, 0.4 * cm))

    # ფუტერი
    story.append(Paragraph(f"<b>აპარატურა:</b> {urinalysis_template['footer']['equipment']}", georgian_style))
    story.append(Paragraph(f"<b>გამოკვლევა შეასრულა:</b> {form_data.get('doctor_name', '')}", georgian_style))
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph("ხელმოწერა: _________________________", georgian_style))

    doc.build(story)
    buffer.seek(0)
    return buffer


@app.route('/')
def ur_form():
    return render_template('form_urinalysis.html', template=urinalysis_template)


@app.route('/generate_urinalysis_doc', methods=['POST'])
def generate_urinalysis_doc():
    form_data = request.form.to_dict()
    doc = create_urinalysis_document(form_data)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"Urinalysis_{form_data.get('last_name', '')}_{datetime.now().strftime('%Y%m%d')}.docx"
    return send_file(
        buf,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename
    )


@app.route('/generate_urinalysis_pdf', methods=['POST'])
def generate_urinalysis_pdf():
    form_data = request.form.to_dict()
    buffer = create_pdf_document(form_data)

    # PDF კონტენტი ავტო-ბეჭდვით
    pdf_content = buffer.getvalue()

    html_wrapper = f'''
    <!DOCTYPE html>
    <html>
    <head>
        <title>Urinalysis Report - Print</title>
        <style>
            body, html {{ margin: 0; padding: 0; height: 100%; overflow: hidden; }}
            iframe {{ width: 100%; height: 100%; border: none; }}
        </style>
    </head>
    <body>
        <iframe id="pdfFrame" src="data:application/pdf;base64,{base64.b64encode(pdf_content).decode()}"></iframe>
        <script>
            window.onload = function() {{
                setTimeout(function() {{
                    var frame = document.getElementById('pdfFrame');
                    frame.contentWindow.focus();
                    frame.contentWindow.print();
                }}, 1000);
            }};
        </script>
    </body>
    </html>
    '''
    return Response(html_wrapper, mimetype='text/html')


if __name__ == '__main__':
    print("=" * 50)
    print("🧪 Urinalysis აპლიკაცია გაშვებულია")
    print("=" * 50)
    print("🌐 გახსენით: http://127.0.0.1:5001")
    print("=" * 50)
    app.run(debug=False, host='127.0.0.1', port=5001)