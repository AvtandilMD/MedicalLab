import time
import webbrowser
import os
from multiprocessing import Process, freeze_support


def get_base_path():
    return os.path.dirname(os.path.abspath(__file__))


# ====== CBC აპლიკაცია (პორტი 5000) ======
def run_cbc():
    os.chdir(get_base_path())
    from flask import Flask, render_template, request, send_file, Response
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io
    from datetime import datetime

    base_path = get_base_path()
    template_folder = os.path.join(base_path, 'templates')
    app = Flask(__name__, template_folder=template_folder)

    cbc_template = {
        "cbc_analysis": [
            {"abbr": "WBC", "parameter": "ლეიკოციტი", "reference_range": "მ. 5.0-10.0; ქ. 5.0-10.0", "unit": "10^9/L"},
            {"abbr": "RBC", "parameter": "ერითროციტი", "reference_range": "მ. 4.5-5.5; ქ. 4.5-5.5", "unit": "10^12/L"},
            {"abbr": "HGB", "parameter": "ჰემოგლობინი", "reference_range": "მ. 140-174; ქ. 120-174", "unit": "g/L"},
            {"abbr": "HCT", "parameter": "ჰემატოკრიტი", "reference_range": "მ. 36-52; ქ. 45-52", "unit": "%"},
            {"abbr": "PLT", "parameter": "თრომბოციტი", "reference_range": "მ. 150-400; ქ. 150-400", "unit": "10^9/L"},
            {"abbr": "RET", "parameter": "რეტიკულოციტი", "reference_range": "მ. 2-10; ქ. 2-10", "unit": "%"},
            {"abbr": "MCV", "parameter": "ერითროც. საშუალო მოცულობა", "reference_range": "მ. 84-96; ქ. 76-96",
             "unit": "FL"},
            {"abbr": "MCH", "parameter": "HGB საშუალო შემცველობა", "reference_range": "მ. 27-32; ქ. 27-32",
             "unit": "pg"},
            {"abbr": "MCHC", "parameter": "HGB საშუალო კონცენტრაცია", "reference_range": "მ. 300-350; ქ. 300-350",
             "unit": "g/l"},
            {"abbr": "RDW", "parameter": "ერითროც. განაწილების ფართი", "reference_range": "მ. 20-42; ქ. 20-42",
             "unit": "%"},
            {"abbr": "MPV", "parameter": "თრომბოც. საშუალო მოცულობა", "reference_range": "მ. 8-15; ქ. 8-15",
             "unit": "FL"},
            {"abbr": "PDW", "parameter": "თრომბოც. განაწილების ფართი", "reference_range": "მ. - ; ქ. -", "unit": "%"},
            {"abbr": "ESR", "parameter": "ერითროც. დალექვის სიჩქარე", "reference_range": "მ. 2-10; ქ. 2-15",
             "unit": "მმ/სთ"}
        ],
        "leukocyte_formula": [
            {"parameter": "მიელოციტი (MIEL %)", "norm": "0%"},
            {"parameter": "მეტამიელოციტი (METAM %)", "norm": "0%"},
            {"parameter": "ჩხირბირთვიანი ნეიტროფილი (Rod NEUT %)", "norm": "0-6%"},
            {"parameter": "სეგმენტბირთვიანი ნეიტროფილი (SEG %)", "norm": "47-72%"},
            {"parameter": "ეოზინოფილი (EO %)", "norm": "0.5-5%"},
            {"parameter": "ბაზოფილი (BASO %)", "norm": "0-1%"},
            {"parameter": "ლიმფოციტი (LYMPH %)", "norm": "19-37%"},
            {"parameter": "მონოციტი (MONO %)", "norm": "3-11%"},
            {"parameter": "პლაზმური უჯრედი (PLAZ %)", "norm": "0.5-1%"}
        ]
    }

    def set_cell_shading(cell, color):
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def create_word_document(form_data):
        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(0.5)
            section.bottom_margin = Cm(0.5)
            section.left_margin = Cm(1.0)
            section.right_margin = Cm(1.0)

        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.paragraph_format.space_after = Pt(0)
        run = header.add_run("PREMIUM MEDI / პრემიუმ მედი")
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 100, 0)

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.paragraph_format.space_after = Pt(2)
        sub.add_run("საოჯახო მედიცინის ცენტრი | ტელ: 558-27-55-51").font.size = Pt(8)

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(4)
        t = title.add_run("BL6 - სისხლის საერთო ანალიზი CBC")
        t.font.size = Pt(10)
        t.font.bold = True

        info = doc.add_paragraph()
        info.paragraph_format.space_after = Pt(4)
        info.add_run("პაციენტი: ").bold = True
        info.add_run(
            f"{form_data.get('first_name', '')} {form_data.get('last_name', '')}, {form_data.get('age', '')} წ.   ")
        info.add_run("თარიღი: ").bold = True
        info.add_run(form_data.get('test_date', ''))
        for r in info.runs:
            r.font.size = Pt(9)

        cbc_title = doc.add_paragraph()
        cbc_title.paragraph_format.space_after = Pt(2)
        cbc_run = cbc_title.add_run("სისხლის საერთო ანალიზი")
        cbc_run.bold = True
        cbc_run.font.size = Pt(9)

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        headers = ['აბრევ.', 'პარამეტრი', 'შედეგი', 'ნორმა', 'ერთ.']
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            set_cell_shading(cell, 'D9E2F3')
            cell.paragraphs[0].runs[0].font.size = Pt(7)
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)

        for item in cbc_template["cbc_analysis"]:
            row = table.add_row()
            row.cells[0].text = item['abbr']
            row.cells[1].text = item['parameter']
            row.cells[2].text = form_data.get(f"cbc_{item['abbr']}", '')
            row.cells[3].text = item['reference_range']
            row.cells[4].text = item['unit']
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(0)
                    for r in p.runs:
                        r.font.size = Pt(7)

        leuko_title = doc.add_paragraph()
        leuko_title.paragraph_format.space_after = Pt(2)
        leuko_title.paragraph_format.space_before = Pt(4)
        leuko_run = leuko_title.add_run("ლეიკოციტარული ფორმულა")
        leuko_run.bold = True
        leuko_run.font.size = Pt(9)

        lt = doc.add_table(rows=1, cols=3)
        lt.style = 'Table Grid'
        lheaders = ['პარამეტრი', 'შედეგი', 'ნორმა']
        for i, h in enumerate(lheaders):
            cell = lt.rows[0].cells[i]
            cell.text = h
            set_cell_shading(cell, 'E2F0D9')
            cell.paragraphs[0].runs[0].font.size = Pt(7)
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)

        for idx, item in enumerate(cbc_template["leukocyte_formula"]):
            row = lt.add_row()
            row.cells[0].text = item['parameter']
            row.cells[1].text = form_data.get(f'leuko_{idx}', '')
            row.cells[2].text = item['norm']
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(0)
                    for r in p.runs:
                        r.font.size = Pt(7)

        morph = doc.add_paragraph()
        morph.paragraph_format.space_after = Pt(0)
        morph.paragraph_format.space_before = Pt(4)
        morph.add_run("ერითროც. მორფოლოგია: ").bold = True
        morph.add_run(form_data.get('erythrocyte_morphology', '') + "   ")
        morph.add_run("ლეიკოც. მორფოლოგია: ").bold = True
        morph.add_run(form_data.get('leukocyte_morphology', ''))
        for r in morph.runs:
            r.font.size = Pt(8)

        footer = doc.add_paragraph()
        footer.paragraph_format.space_after = Pt(0)
        footer.paragraph_format.space_before = Pt(6)
        footer.add_run("გამოკვლევა შეასრულა: ").bold = True
        footer.add_run(form_data.get('doctor_name', '') + "          ")
        footer.add_run("ხელმოწერა: _____________")
        for r in footer.runs:
            r.font.size = Pt(8)

        return doc

    def create_print_html(form_data):
        first_name = form_data.get('first_name', '')
        last_name = form_data.get('last_name', '')
        age = form_data.get('age', '')
        test_date = form_data.get('test_date', '')
        erythrocyte_morphology = form_data.get('erythrocyte_morphology', '')
        leukocyte_morphology = form_data.get('leukocyte_morphology', '')
        doctor_name = form_data.get('doctor_name', '')

        html = f'''<!DOCTYPE html>
<html><head><meta charset="UTF-8"><title>CBC Report</title>
<style>
@page {{ size: A4; margin: 10mm; }}
body {{ font-family: Arial, sans-serif; padding: 10px; font-size: 11px; }}
h1 {{ color: green; text-align: center; font-size: 14px; margin: 5px 0; }}
h2 {{ text-align: center; font-size: 12px; margin: 5px 0; }}
h3 {{ font-size: 10px; margin: 8px 0 4px 0; }}
p {{ margin: 3px 0; }}
table {{ width: 100%; border-collapse: collapse; margin: 5px 0; }}
th, td {{ border: 1px solid #ddd; padding: 4px; text-align: left; font-size: 9px; }}
th {{ background: #D9E2F3; }}
.leuko th {{ background: #E2F0D9; }}
</style></head><body>
<h1>PREMIUM MEDI / პრემიუმ მედი</h1>
<p style="text-align:center; font-size: 9px;">საოჯახო მედიცინის ცენტრი | ტელ: 558-27-55-51</p>
<h2>BL6 - სისხლის საერთო ანალიზი CBC</h2>
<p><b>პაციენტი:</b> {first_name} {last_name}, {age} წ. &nbsp;&nbsp; <b>თარიღი:</b> {test_date}</p>
<h3>სისხლის საერთო ანალიზი</h3>
<table><tr><th>აბრევ.</th><th>პარამეტრი</th><th>შედეგი</th><th>ნორმა</th><th>ერთ.</th></tr>'''

        for item in cbc_template["cbc_analysis"]:
            abbr = item['abbr']
            result = form_data.get(f'cbc_{abbr}', '')
            html += f"<tr><td>{abbr}</td><td>{item['parameter']}</td><td><b>{result}</b></td><td>{item['reference_range']}</td><td>{item['unit']}</td></tr>"

        html += '</table><h3>ლეიკოციტარული ფორმულა</h3><table class="leuko"><tr><th>პარამეტრი</th><th>შედეგი</th><th>ნორმა</th></tr>'

        for idx, item in enumerate(cbc_template["leukocyte_formula"]):
            result = form_data.get(f'leuko_{idx}', '')
            html += f"<tr><td>{item['parameter']}</td><td><b>{result}</b></td><td>{item['norm']}</td></tr>"

        html += f'''</table>
<p><b>ერითროც. მორფოლოგია:</b> {erythrocyte_morphology} &nbsp;&nbsp; <b>ლეიკოც. მორფოლოგია:</b> {leukocyte_morphology}</p>
<p><b>გამოკვლევა შეასრულა:</b> {doctor_name} &nbsp;&nbsp;&nbsp;&nbsp; <b>ხელმოწერა:</b> _____________</p>
<script>window.onload = function() {{ setTimeout(function() {{ window.print(); }}, 500); }}</script>
</body></html>'''
        return html

    @app.route('/')
    def index():
        return render_template('form_cbc.html', template=cbc_template)

    @app.route('/cbc/print', methods=['POST'])
    @app.route('/generate_cbc_pdf', methods=['POST'])
    def generate_pdf():
        form_data = request.form.to_dict()
        html = create_print_html(form_data)
        return Response(html, mimetype='text/html')

    @app.route('/generate_cbc_doc', methods=['POST'])
    def generate_doc():
        form_data = request.form.to_dict()
        doc = create_word_document(form_data)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        filename = f"CBC_{form_data.get('last_name', '')}_{datetime.now().strftime('%Y%m%d')}.docx"
        return send_file(buffer, as_attachment=True, download_name=filename,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    print("🩸 CBC სერვერი: http://127.0.0.1:5000")
    app.run(debug=False, host='127.0.0.1', port=5000, use_reloader=False)


# ====== URINE აპლიკაცია (პორტი 5001) ======
def run_urine():
    os.chdir(get_base_path())
    from flask import Flask, render_template, request, send_file, Response
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io
    from datetime import datetime

    base_path = get_base_path()
    template_folder = os.path.join(base_path, 'templates')
    app = Flask(__name__, template_folder=template_folder)

    urinalysis_template = {
        "header": {
            "subtitle": "საოჯახო მედიცინის ცენტრი",
            "phones": ["558-27-55-51", "577-03-97-70"]
        },
        "test_info": {"code": "UR.7", "name": "შარდის საერთო ანალიზი"},
        "physico_chemical": [
            {"abbr": "", "parameter": "რაოდენობა", "norm": "", "unit": "მლ"},
            {"abbr": "", "parameter": "ფერი", "norm": "ჩალისფერი", "unit": ""},
            {"abbr": "", "parameter": "გამჭვირვალობა", "norm": "გამჭვირვალე", "unit": ""},
            {"abbr": "SG", "parameter": "ხვედრითი წონა", "norm": "1.005-1.030", "unit": ""},
            {"abbr": "PH", "parameter": "რეაქცია", "norm": "5.0-8.0", "unit": ""},
            {"abbr": "PRO", "parameter": "ცილა", "norm": "0", "unit": "g/l"},
            {"abbr": "GLU", "parameter": "გლუკოზა", "norm": "0", "unit": "mmol/l"},
            {"abbr": "KET", "parameter": "კეტონები", "norm": "0", "unit": "mmol/l"},
            {"abbr": "UBG", "parameter": "ურობილინოგენი", "norm": "3.4-17.0", "unit": "µmol/l"},
            {"abbr": "BIL", "parameter": "ბილირუბინი", "norm": "0", "unit": "µmol/l"},
            {"abbr": "NIT", "parameter": "ნიტრატები", "norm": "NEG", "unit": ""},
            {"abbr": "LEU", "parameter": "ლეიკოციტები", "norm": "-", "unit": "Leu/µL"},
            {"abbr": "BLD", "parameter": "ერითროციტები", "norm": "-", "unit": "Ery/µL"}
        ],
        "microscopy": {
            "epithelium": [
                {"key": "squamous", "label": "ბრტყელი"},
                {"key": "transitional", "label": "გარდამავალი"},
                {"key": "renal", "label": "თირკმლის"}
            ],
            "cylinders": [
                {"key": "hyaline", "label": "ჰიალინური"},
                {"key": "granular", "label": "მარცვლოვანი"},
                {"key": "waxy", "label": "ცვილისებური"}
            ],
            "others": [
                {"key": "mucus", "parameter": "ლორწო"},
                {"key": "salts", "parameter": "მარილები"},
                {"key": "bacteria", "parameter": "ბაქტერიები"},
                {"key": "fungi", "parameter": "სოკო"}
            ]
        },
        "footer": {"equipment": "SIEMENS CLINITEK Status+"}
    }

    def set_cell_shading(cell, color):
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def create_urinalysis_document(form_data):
        doc = Document()
        for s in doc.sections:
            s.top_margin = Cm(0.5)
            s.bottom_margin = Cm(0.5)
            s.left_margin = Cm(1.0)
            s.right_margin = Cm(1.0)

        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.paragraph_format.space_after = Pt(0)
        r1 = header.add_run("PREMIUM MEDI / პრემიუმ მედი")
        r1.font.size = Pt(12)
        r1.font.bold = True
        r1.font.color.rgb = RGBColor(0, 100, 0)

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.paragraph_format.space_after = Pt(2)
        phones = ', '.join(urinalysis_template['header']['phones'])
        sub.add_run(f"{urinalysis_template['header']['subtitle']} | ტელ: {phones}").font.size = Pt(8)

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(4)
        t = title.add_run(f"{urinalysis_template['test_info']['code']} - {urinalysis_template['test_info']['name']}")
        t.font.size = Pt(10)
        t.font.bold = True

        info = doc.add_paragraph()
        info.paragraph_format.space_after = Pt(4)
        info.add_run("პაციენტი: ").bold = True
        info.add_run(
            f"{form_data.get('first_name', '')} {form_data.get('last_name', '')}, {form_data.get('age', '')} წ.   ")
        info.add_run("თარიღი: ").bold = True
        info.add_run(form_data.get('test_date', ''))
        for r in info.runs:
            r.font.size = Pt(9)

        phys_title = doc.add_paragraph()
        phys_title.paragraph_format.space_after = Pt(2)
        phys_run = phys_title.add_run("ფიზიკო-ქიმიური თვისებები")
        phys_run.bold = True
        phys_run.font.size = Pt(9)

        t1 = doc.add_table(rows=1, cols=5)
        t1.style = 'Table Grid'
        h1 = ['აბრევ.', 'პარამეტრი', 'შედეგი', 'ნორმა', 'ერთ.']
        for i, h in enumerate(h1):
            cell = t1.rows[0].cells[i]
            cell.text = h
            set_cell_shading(cell, 'FFF2CC')
            cell.paragraphs[0].runs[0].font.size = Pt(7)
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)

        for idx, item in enumerate(urinalysis_template["physico_chemical"]):
            row = t1.add_row()
            row.cells[0].text = item['abbr']
            row.cells[1].text = item['parameter']
            row.cells[2].text = form_data.get(f'phys_{idx}', '')
            row.cells[3].text = item['norm']
            row.cells[4].text = item['unit']
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(0)
                    for r in p.runs:
                        r.font.size = Pt(7)

        micro_title = doc.add_paragraph()
        micro_title.paragraph_format.space_after = Pt(2)
        micro_title.paragraph_format.space_before = Pt(4)
        micro_run = micro_title.add_run("მიკროსკოპია")
        micro_run.bold = True
        micro_run.font.size = Pt(9)

        mt = doc.add_table(rows=1, cols=4)
        mt.style = 'Table Grid'
        mh = ['ეპითელიუმი', 'შედეგი', 'ცილინდრები', 'შედეგი']
        for i, h in enumerate(mh):
            cell = mt.rows[0].cells[i]
            cell.text = h
            set_cell_shading(cell, 'E2EFDA')
            cell.paragraphs[0].runs[0].font.size = Pt(7)
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)

        epi = urinalysis_template["microscopy"]["epithelium"]
        cyl = urinalysis_template["microscopy"]["cylinders"]
        for i in range(max(len(epi), len(cyl))):
            row = mt.add_row()
            if i < len(epi):
                row.cells[0].text = epi[i]['label']
                row.cells[1].text = form_data.get(f"epi_{epi[i]['key']}", '')
            if i < len(cyl):
                row.cells[2].text = cyl[i]['label']
                row.cells[3].text = form_data.get(f"cyl_{cyl[i]['key']}", '')
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(0)
                    for r in p.runs:
                        r.font.size = Pt(7)

        other_title = doc.add_paragraph()
        other_title.paragraph_format.space_after = Pt(2)
        other_title.paragraph_format.space_before = Pt(4)
        other_run = other_title.add_run("სხვა მონაცემები")
        other_run.bold = True
        other_run.font.size = Pt(9)

        ot = doc.add_table(rows=1, cols=4)
        ot.style = 'Table Grid'
        oh = ['პარამეტრი', 'შედეგი', 'პარამეტრი', 'შედეგი']
        for i, h in enumerate(oh):
            cell = ot.rows[0].cells[i]
            cell.text = h
            set_cell_shading(cell, 'DDEBF7')
            cell.paragraphs[0].runs[0].font.size = Pt(7)
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)

        others = urinalysis_template["microscopy"]["others"]
        for i in range(0, len(others), 2):
            row = ot.add_row()
            row.cells[0].text = others[i]['parameter']
            row.cells[1].text = form_data.get(f"other_{others[i]['key']}", '')
            if i + 1 < len(others):
                row.cells[2].text = others[i + 1]['parameter']
                row.cells[3].text = form_data.get(f"other_{others[i + 1]['key']}", '')
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(0)
                    for r in p.runs:
                        r.font.size = Pt(7)

        footer = doc.add_paragraph()
        footer.paragraph_format.space_before = Pt(6)
        footer.add_run("აპარატურა: ").bold = True
        footer.add_run(urinalysis_template["footer"]["equipment"] + "   ")
        footer.add_run("შეასრულა: ").bold = True
        footer.add_run(form_data.get('doctor_name', '') + "   ")
        footer.add_run("ხელმოწერა: _____________")
        for r in footer.runs:
            r.font.size = Pt(8)

        return doc

    def create_print_html(form_data):
        phones = ', '.join(urinalysis_template['header']['phones'])
        html = f'''<!DOCTYPE html>
<html><head><meta charset="UTF-8"><title>Urinalysis Report</title>
<style>
@page {{ size: A4; margin: 10mm; }}
body {{ font-family: Arial, sans-serif; padding: 10px; font-size: 11px; }}
h1 {{ color: green; text-align: center; font-size: 14px; margin: 5px 0; }}
h2 {{ text-align: center; font-size: 12px; margin: 5px 0; }}
h3 {{ font-size: 10px; margin: 8px 0 4px 0; }}
p {{ margin: 3px 0; }}
table {{ width: 100%; border-collapse: collapse; margin: 5px 0; }}
th, td {{ border: 1px solid #ddd; padding: 4px; text-align: left; font-size: 9px; }}
th {{ background: #FFF2CC; }}
.micro th {{ background: #E2EFDA; }}
.other th {{ background: #DDEBF7; }}
</style></head><body>
<h1>PREMIUM MEDI / პრემიუმ მედი</h1>
<p style="text-align:center; font-size: 9px;">{urinalysis_template['header']['subtitle']} | ტელ: {phones}</p>
<h2>{urinalysis_template['test_info']['code']} - {urinalysis_template['test_info']['name']}</h2>
<p><b>პაციენტი:</b> {form_data.get('first_name', '')} {form_data.get('last_name', '')}, {form_data.get('age', '')} წ. &nbsp;&nbsp; <b>თარიღი:</b> {form_data.get('test_date', '')}</p>
<h3>ფიზიკო-ქიმიური თვისებები</h3>
<table><tr><th>აბრევ.</th><th>პარამეტრი</th><th>შედეგი</th><th>ნორმა</th><th>ერთ.</th></tr>'''

        for idx, item in enumerate(urinalysis_template["physico_chemical"]):
            result = form_data.get(f'phys_{idx}', '')
            html += f"<tr><td>{item['abbr']}</td><td>{item['parameter']}</td><td><b>{result}</b></td><td>{item['norm']}</td><td>{item['unit']}</td></tr>"

        html += '</table><h3>მიკროსკოპია</h3><table class="micro"><tr><th>ეპითელიუმი</th><th>შედეგი</th><th>ცილინდრები</th><th>შედეგი</th></tr>'

        epi = urinalysis_template["microscopy"]["epithelium"]
        cyl = urinalysis_template["microscopy"]["cylinders"]
        for i in range(max(len(epi), len(cyl))):
            e_label = epi[i]['label'] if i < len(epi) else ''
            e_val = form_data.get(f"epi_{epi[i]['key']}", '') if i < len(epi) else ''
            c_label = cyl[i]['label'] if i < len(cyl) else ''
            c_val = form_data.get(f"cyl_{cyl[i]['key']}", '') if i < len(cyl) else ''
            html += f"<tr><td>{e_label}</td><td><b>{e_val}</b></td><td>{c_label}</td><td><b>{c_val}</b></td></tr>"

        html += '</table><h3>სხვა მონაცემები</h3><table class="other"><tr><th>პარამეტრი</th><th>შედეგი</th><th>პარამეტრი</th><th>შედეგი</th></tr>'

        others = urinalysis_template["microscopy"]["others"]
        for i in range(0, len(others), 2):
            p1 = others[i]['parameter']
            v1 = form_data.get(f"other_{others[i]['key']}", '')
            p2 = others[i + 1]['parameter'] if i + 1 < len(others) else ''
            v2 = form_data.get(f"other_{others[i + 1]['key']}", '') if i + 1 < len(others) else ''
            html += f"<tr><td>{p1}</td><td><b>{v1}</b></td><td>{p2}</td><td><b>{v2}</b></td></tr>"

        equipment = urinalysis_template["footer"]["equipment"]
        doctor = form_data.get('doctor_name', '')
        html += f'''</table>
<p><b>აპარატურა:</b> {equipment} &nbsp;&nbsp; <b>შეასრულა:</b> {doctor} &nbsp;&nbsp; <b>ხელმოწერა:</b> _____________</p>
<script>window.onload = function() {{ setTimeout(function() {{ window.print(); }}, 500); }}</script>
</body></html>'''
        return html

    @app.route('/')
    def ur_form():
        return render_template('form_urinalysis.html', template=urinalysis_template)

    @app.route('/urine/print', methods=['POST'])
    @app.route('/generate_urinalysis_pdf', methods=['POST'])
    def generate_urinalysis_pdf():
        form_data = request.form.to_dict()
        html = create_print_html(form_data)
        return Response(html, mimetype='text/html')

    @app.route('/generate_urinalysis_doc', methods=['POST'])
    def generate_urinalysis_doc():
        form_data = request.form.to_dict()
        doc = create_urinalysis_document(form_data)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        filename = f"Urinalysis_{form_data.get('last_name', '')}_{datetime.now().strftime('%Y%m%d')}.docx"
        return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True, download_name=filename)

    print("🧪 Urine სერვერი: http://127.0.0.1:5001")
    app.run(debug=False, host='127.0.0.1', port=5001, use_reloader=False)


# ====== CRP აპლიკაცია (პორტი 5002) ======
def run_crp():
    os.chdir(get_base_path())
    from flask import Flask, render_template, request, send_file, Response
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io
    from datetime import datetime

    base_path = get_base_path()
    template_folder = os.path.join(base_path, 'templates')
    app = Flask(__name__, template_folder=template_folder)

    crp_template = {
        "clinic_info": {
            "name": "პრემიუმ მედი",
            "description": "საოჯახო მედიცინის ცენტრი",
            "phones": ["558-27-55-51", "577-03-97-70"]
        },
        "test_details": {
            "title_ge": "მაღალი მგრძნობელობის C-რეაქტიული ცილა (BL.7.9.1)",
            "title_en": "High-Sensitivity C-Reactive Protein"
        },
        "test_results": [
            {"code": "CRP", "parameter": "C-რეაქტიული ცილა", "reference_range": "0-10", "unit": "mg/L (მგ/ლ)"},
            {"code": "hsCRP", "parameter": "მაღალი მგრძნობელობის C-რეაქტიული ცილა", "reference_range": "0-1",
             "unit": "mg/L (მგ/ლ)"}
        ]
    }

    def set_cell_shading(cell, color):
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def create_crp_document(form_data):
        doc = Document()
        for s in doc.sections:
            s.top_margin = Cm(1.5)
            s.bottom_margin = Cm(1.5)
            s.left_margin = Cm(2.0)
            s.right_margin = Cm(2.0)

        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.paragraph_format.space_after = Pt(6)
        r1 = header.add_run("PREMIUM MEDI / პრემიუმ მედი")
        r1.font.size = Pt(16)
        r1.font.bold = True
        r1.font.color.rgb = RGBColor(0, 100, 0)

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.paragraph_format.space_after = Pt(12)
        phones = ', '.join(crp_template['clinic_info']['phones'])
        sub.add_run(f"{crp_template['clinic_info']['description']} | ტელ: {phones}").font.size = Pt(10)

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(16)
        t = title.add_run(crp_template['test_details']['title_ge'])
        t.font.size = Pt(14)
        t.font.bold = True

        info = doc.add_paragraph()
        info.paragraph_format.space_after = Pt(16)
        info.add_run("პაციენტი: ").bold = True
        info.add_run(
            f"{form_data.get('first_name', '')} {form_data.get('last_name', '')}, {form_data.get('age', '')} წ.          ")
        info.add_run("თარიღი: ").bold = True
        info.add_run(form_data.get('test_date', ''))
        for r in info.runs:
            r.font.size = Pt(11)

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        headers = ['კოდი', 'პარამეტრი', 'შედეგი', 'ნორმა', 'ერთეული']
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            set_cell_shading(cell, 'E8DAEF')
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].runs[0].font.bold = True

        for item in crp_template["test_results"]:
            row = table.add_row()
            row.cells[0].text = item['code']
            row.cells[1].text = item['parameter']
            row.cells[2].text = form_data.get(f"res_{item['code']}", '')
            row.cells[3].text = item['reference_range']
            row.cells[4].text = item['unit']
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10)

        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.paragraph_format.space_before = Pt(24)
        footer.add_run("გამოკვლევა შეასრულა: ").bold = True
        footer.add_run(form_data.get('doctor_name', '') + "                    ")
        footer.add_run("ხელმოწერა: _________________________")
        for r in footer.runs:
            r.font.size = Pt(10)

        return doc

    def create_print_html(form_data):
        phones = ', '.join(crp_template['clinic_info']['phones'])
        html = f'''<!DOCTYPE html>
<html><head><meta charset="UTF-8"><title>CRP Report</title>
<style>
@page {{ size: A4; margin: 20mm; }}
body {{ font-family: Arial, sans-serif; padding: 20px; }}
h1 {{ color: green; text-align: center; font-size: 18px; margin: 10px 0; }}
h2 {{ text-align: center; font-size: 16px; margin: 15px 0; color: #8e44ad; }}
p {{ margin: 8px 0; font-size: 12px; }}
table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
th, td {{ border: 1px solid #ddd; padding: 12px; text-align: left; font-size: 12px; }}
th {{ background: #E8DAEF; font-weight: bold; }}
.result {{ font-weight: bold; font-size: 14px; }}
</style></head><body>
<h1>PREMIUM MEDI / პრემიუმ მედი</h1>
<p style="text-align:center;">{crp_template['clinic_info']['description']} | ტელ: {phones}</p>
<h2>{crp_template['test_details']['title_ge']}</h2>
<p><b>პაციენტი:</b> {form_data.get('first_name', '')} {form_data.get('last_name', '')}, {form_data.get('age', '')} წ. &nbsp;&nbsp;&nbsp;&nbsp;&nbsp; <b>თარიღი:</b> {form_data.get('test_date', '')}</p>
<table>
<tr><th>კოდი</th><th>პარამეტრი</th><th>შედეგი</th><th>ნორმა</th><th>ერთეული</th></tr>'''

        for item in crp_template["test_results"]:
            result = form_data.get(f"res_{item['code']}", '')
            html += f"<tr><td><b>{item['code']}</b></td><td>{item['parameter']}</td><td class='result'>{result}</td><td>{item['reference_range']}</td><td>{item['unit']}</td></tr>"

        doctor = form_data.get('doctor_name', '')
        html += f'''</table>
<br><br>
<p><b>გამოკვლევა შეასრულა:</b> {doctor} &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; <b>ხელმოწერა:</b> _________________________</p>
<script>window.onload = function() {{ setTimeout(function() {{ window.print(); }}, 500); }}</script>
</body></html>'''
        return html

    @app.route('/')
    def index():
        return render_template('form_crp.html', template=crp_template)

    @app.route('/crp/print', methods=['POST'])
    @app.route('/generate_crp_pdf', methods=['POST'])
    def generate_crp_pdf():
        form_data = request.form.to_dict()
        html = create_print_html(form_data)
        return Response(html, mimetype='text/html')

    @app.route('/generate_crp_doc', methods=['POST'])
    def generate_crp_doc():
        form_data = request.form.to_dict()
        doc = create_crp_document(form_data)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        filename = f"CRP_{form_data.get('last_name', '')}_{datetime.now().strftime('%Y%m%d')}.docx"
        return send_file(buffer, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True, download_name=filename)

    print("🧬 CRP სერვერი: http://127.0.0.1:5002")
    app.run(debug=False, host='127.0.0.1', port=5002, use_reloader=False)


# ====== LIPID აპლიკაცია (პორტი 5003) ======
def run_lipid():
    os.chdir(get_base_path())
    from flask import Flask, render_template, request, send_file, Response
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io
    from datetime import datetime

    base_path = get_base_path()
    template_folder = os.path.join(base_path, 'templates')
    app = Flask(__name__, template_folder=template_folder)

    lipid_template = {
        "clinic_info": {
            "description": "საოჯახო მედიცინის ცენტრი",
            "phones": ["558-27-55-51", "577-03-97-70"]
        },
        "test_details": {
            "title": "ანალიზის პასუხები",
            "subtitle": "ლიპიდური სპექტრი"
        },
        "test_results": [
            {
                "code": "TC",
                "name": "საერთო ქოლესტერინი TC",
                "desirable": "<5.2 მმოლ/ლ\n(200 მგ/დლ)",
                "borderline": "5.2-6.2 მმოლ/ლ\n(200-240 მგ/დლ)",
                "high": ">6.2 მმოლ/ლ\n(240 მგ/დლ)"
            },
            {
                "code": "HDL",
                "name": "მაღალი სიმკვრივის ქოლესტერინი HDL",
                "desirable": "≥1.5 მმოლ/ლ\n(60 მგ/დლ)",
                "borderline": "კაცები: 1.5 -1.0 მმოლ/ლ\n(60-40 მგ/დლ)\nქალები: 1.5 - 1.3 მმოლ/ლ\n(60-50 მგ/დლ)",
                "high": "კაცები: <1.0 მმოლ/ლ\n(40 მგ/დლ)\nქალები: <1.3 მმოლ/ლ\n(50 მგ/დლ)"
            },
            {
                "code": "TG",
                "name": "ტრიგლიცერიდები TG",
                "desirable": "<1.7 მმოლ/ლ\n(150 მგ/დლ)",
                "borderline": "1.7 -2.3 მმოლ/ლ\n(150-200 მგ/დლ)",
                "high": ">2.3 მმოლ/ლ\n(200 მგ/დლ)"
            },
            {
                "code": "TC_HDL",
                "name": "საერთო ქოლესტერინი/მაღალი სიმკვრივის ქოლესტერინი TC/HDL",
                "desirable": "<4.5",
                "borderline": "4.5-5.0",
                "high": ">5.0"
            },
            {
                "code": "LDL",
                "name": "დაბალი სიმკვრივის ქოლესტერინი LDL",
                "desirable": "<3.4 მმოლ/ლ\n(130 მგ/დლ)",
                "borderline": "3.4-4.1 მმოლ/ლ\n(130-160 მგ/დლ)",
                "high": ">4.1 მმოლ/ლ\n(160 მგ/დლ)"
            }
        ]
    }

    def set_cell_shading(cell, color):
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def create_lipid_document(form_data):
        doc = Document()
        for s in doc.sections:
            s.top_margin = Cm(1.5)
            s.bottom_margin = Cm(1.5)
            s.left_margin = Cm(1.5)
            s.right_margin = Cm(1.5)

        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.paragraph_format.space_after = Pt(6)
        r1 = header.add_run("PREMIUM MEDI / პრემიუმ მედი")
        r1.font.size = Pt(16)
        r1.font.bold = True
        r1.font.color.rgb = RGBColor(0, 100, 0)

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.paragraph_format.space_after = Pt(12)
        phones = ', '.join(lipid_template['clinic_info']['phones'])
        sub.add_run(f"{lipid_template['clinic_info']['description']} | ტელ: {phones}").font.size = Pt(10)

        info_table = doc.add_table(rows=3, cols=2)
        info_table.style = 'Table Grid'
        info_table.rows[0].cells[
            0].text = f"პაციენტი: {form_data.get('first_name', '')} {form_data.get('last_name', '')}"
        info_table.rows[0].cells[1].text = f"სამედიცინო დაწ.: {form_data.get('facility', 'PREMIUM MEDI')}"
        info_table.rows[1].cells[0].text = f"პირადი ნომერი: {form_data.get('personal_id', '')}"
        info_table.rows[1].cells[1].text = f"ექიმი: {form_data.get('doctor_name', '')}"
        info_table.rows[2].cells[0].text = f"სქესი და ასაკი: {form_data.get('sex', '')} / {form_data.get('age', '')} წ."
        info_table.rows[2].cells[1].text = f"თარიღი: {form_data.get('test_date', '')}"

        for row in info_table.rows:
            for cell in row.cells:
                set_cell_shading(cell, 'F2F2F2')
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(4)
                    p.paragraph_format.space_before = Pt(4)
                    for r in p.runs:
                        r.font.size = Pt(10)

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Pt(16)
        title.paragraph_format.space_after = Pt(4)
        t = title.add_run(lipid_template['test_details']['title'])
        t.font.size = Pt(16)
        t.font.bold = True

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(12)
        st = subtitle.add_run(lipid_template['test_details']['subtitle'])
        st.font.size = Pt(12)

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        headers = ['დასახელება', 'სასურველი', 'ზღვარზე მაღალი', 'მაღალი', 'შედეგი']
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            set_cell_shading(cell, 'E0E0E0')
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].runs[0].font.bold = True

        for item in lipid_template["test_results"]:
            row = table.add_row()
            row.cells[0].text = item['name']
            row.cells[1].text = item['desirable']
            row.cells[2].text = item['borderline']
            row.cells[3].text = item['high']

            res_cell = row.cells[4]
            res_cell.text = form_data.get(f"res_{item['code']}", '')
            res_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                res_cell.paragraphs[0].runs[0].font.bold = True
            except:
                pass

            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)
                    for r in p.runs:
                        r.font.size = Pt(9)

        return doc

    def create_print_html(form_data):
        html = f'''<!DOCTYPE html>
<html><head><meta charset="UTF-8"><title>Lipid Profile Report</title>
<style>
@page {{ size: A4; margin: 15mm; }}
body {{ font-family: Arial, sans-serif; padding: 10px; font-size: 12px; }}
h1 {{ color: green; text-align: center; font-size: 18px; margin: 5px 0; }}
h2 {{ text-align: center; font-size: 20px; margin: 15px 0 5px 0; }}
h3 {{ text-align: center; font-size: 14px; font-weight: normal; margin: 0 0 15px 0; }}
.info-table {{ width: 100%; border-collapse: collapse; margin-bottom: 20px; background: #f9f9f9; }}
.info-table td {{ padding: 10px; border: 1px solid #ddd; }}
.result-table {{ width: 100%; border-collapse: collapse; margin: 10px 0; }}
.result-table th, .result-table td {{ border: 1px solid #aaa; padding: 8px; text-align: center; vertical-align: middle; }}
.result-table th {{ background: #e0e0e0; font-weight: bold; font-size: 13px; }}
.result-table td:first-child {{ text-align: left; font-weight: bold; width: 25%; }}
.val {{ font-weight: bold; font-size: 14px; }}
</style></head><body>
<h1>PREMIUM MEDI / პრემიუმ მედი</h1>

<table class="info-table">
    <tr>
        <td><b>პაციენტი:</b> {form_data.get('first_name', '')} {form_data.get('last_name', '')}</td>
        <td><b>სამედიცინო დაწ.:</b> {form_data.get('facility', 'PREMIUM MEDI')}</td>
    </tr>
    <tr>
        <td><b>პირადი ნომერი:</b> {form_data.get('personal_id', '')}</td>
        <td><b>ექიმი:</b> {form_data.get('doctor_name', '')}</td>
    </tr>
    <tr>
        <td><b>სქესი და ასაკი:</b> {form_data.get('sex', '')} / {form_data.get('age', '')} წ.</td>
        <td><b>თარიღი:</b> {form_data.get('test_date', '')}</td>
    </tr>
</table>

<h2>{lipid_template['test_details']['title']}</h2>
<h3>{lipid_template['test_details']['subtitle']}</h3>

<table class="result-table">
<tr><th>დასახელება</th><th>სასურველი</th><th>ზღვარზე მაღალი</th><th>მაღალი</th><th>შედეგი</th></tr>'''

        for item in lipid_template["test_results"]:
            result = form_data.get(f"res_{item['code']}", '')
            desirable = item['desirable'].replace('\n', '<br>')
            borderline = item['borderline'].replace('\n', '<br>')
            high = item['high'].replace('\n', '<br>')
            html += f"<tr><td>{item['name']}</td><td>{desirable}</td><td>{borderline}</td><td>{high}</td><td class='val'>{result}</td></tr>"

        html += '''</table>
<script>window.onload = function() { setTimeout(function() { window.print(); }, 500); }</script>
</body></html>'''
        return html

    @app.route('/')
    def index():
        return render_template('form_lipid.html', template=lipid_template)

    @app.route('/generate_lipid_doc', methods=['POST'])
    def generate_lipid_doc():
        form_data = request.form.to_dict()
        doc = create_lipid_document(form_data)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        filename = f"Lipid_{form_data.get('last_name', '')}_{datetime.now().strftime('%Y%m%d')}.docx"
        return send_file(buffer, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True, download_name=filename)

    @app.route('/lipid/print', methods=['POST'])
    @app.route('/generate_lipid_pdf', methods=['POST'])
    def generate_lipid_pdf():
        form_data = request.form.to_dict()
        html = create_print_html(form_data)
        return Response(html, mimetype='text/html')

    print("🩸 Lipid სერვერი: http://127.0.0.1:5003")
    app.run(debug=True, host='127.0.0.1', port=5003, use_reloader=False)


# ====== TROPONIN აპლიკაცია (პორტი 5004) ======
def run_trop():
    os.chdir(get_base_path())
    from flask import Flask, render_template, request, send_file, Response
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io
    from datetime import datetime

    base_path = get_base_path()
    template_folder = os.path.join(base_path, 'templates')
    app = Flask(__name__, template_folder=template_folder)

    TROPONIN_TEMPLATE = {
        "document_info": {"clinic_description": "საოჯახო მედიცინის ცენტრი",
                          "contact": "ტელ: 558-27-55-51, 577-03-97-70"},
        "test_info": {
            "title": "ტროპონინის ტესტი (BL.7.8)",
            "results_table": [{"code": "BL.7.8", "parameter": "ტროპონინი", "reference_range": "უარყოფითი"}],
        },
        "footer_note": {"equipment": "გამოკვლევა ჩატარდა ანალიზატორ Firance FS-113 _ზე"},
    }

    def set_cell_shading(cell, color):
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def create_troponin_document(fd):
        doc = Document()
        for s in doc.sections:
            s.top_margin = Cm(1.5)
            s.bottom_margin = Cm(1.5)
            s.left_margin = Cm(2.0)
            s.right_margin = Cm(2.0)

        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = h.add_run("PREMIUM MEDI / პრემიუმ მედი")
        r.font.size = Pt(16)
        r.bold = True
        r.font.color.rgb = RGBColor(0, 100, 0)

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.add_run(
            f"{TROPONIN_TEMPLATE['document_info']['clinic_description']} | {TROPONIN_TEMPLATE['document_info']['contact']}").font.size = Pt(
            10)

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Pt(16)
        title.paragraph_format.space_after = Pt(16)
        tr = title.add_run(TROPONIN_TEMPLATE["test_info"]["title"])
        tr.font.size = Pt(14)
        tr.bold = True

        info = doc.add_paragraph()
        info.add_run("პაციენტი: ").bold = True
        info.add_run(f"{fd.get('first_name', '')} {fd.get('last_name', '')}, {fd.get('age', '')} წ.   ")
        info.add_run("თარიღი: ").bold = True
        info.add_run(fd.get("test_date", ""))

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = ["კოდი", "პარამეტრი", "შედეგი", "ნორმა"]
        for i, htxt in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = htxt
            set_cell_shading(cell, "FDEBD0")
            cell.paragraphs[0].runs[0].bold = True

        row = table.add_row()
        row.cells[0].text = "BL.7.8"
        row.cells[1].text = "ტროპონინი"
        row.cells[2].text = fd.get("result_value", "")
        row.cells[3].text = "უარყოფითი"

        eq = doc.add_paragraph()
        eq.paragraph_format.space_before = Pt(16)
        eq.add_run("აპარატურა: ").bold = True
        eq.add_run(TROPONIN_TEMPLATE["footer_note"]["equipment"])

        footer = doc.add_paragraph()
        footer.add_run("შეასრულა: ").bold = True
        footer.add_run(fd.get("doctor_name", "") + "                    ")
        footer.add_run("ხელმოწერა: _________________________")

        return doc

    def create_print_html(fd):
        html = f'''<!DOCTYPE html><html><head><meta charset="UTF-8"><title>Troponin</title>
        <style>
        @page{{size:A4;margin:15mm}}
        body{{font-family:Arial,sans-serif;padding:15px;font-size:12px}}
        h1{{color:green;text-align:center;font-size:18px;margin:4px 0}}
        h2{{text-align:center;font-size:16px;margin:15px 0;color:#d35400}}
        p{{margin:6px 0}}
        table{{width:100%;border-collapse:collapse;margin:15px 0}}
        th,td{{border:1px solid #ddd;padding:10px;text-align:left;font-size:12px}}
        th{{background:#FDEBD0}}
        </style></head><body>
        <h1>PREMIUM MEDI / პრემიუმ მედი</h1>
        <p style="text-align:center">{TROPONIN_TEMPLATE['document_info']['clinic_description']} | {TROPONIN_TEMPLATE['document_info']['contact']}</p>
        <h2>{TROPONIN_TEMPLATE['test_info']['title']}</h2>
        <p><b>პაციენტი:</b> {fd.get('first_name', '')} {fd.get('last_name', '')}, {fd.get('age', '')} წ. &nbsp;&nbsp;&nbsp;&nbsp; <b>თარიღი:</b> {fd.get('test_date', '')}</p>
        <table><tr><th>კოდი</th><th>პარამეტრი</th><th>შედეგი</th><th>ნორმა</th></tr>
        <tr><td><b>BL.7.8</b></td><td>ტროპონინი</td><td><b>{fd.get('result_value', '')}</b></td><td>უარყოფითი</td></tr>
        </table>
        <br>
        <p><b>აპარატურა:</b> {TROPONIN_TEMPLATE['footer_note']['equipment']}</p>
        <p><b>შეასრულა:</b> {fd.get('doctor_name', '')} &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp; <b>ხელმოწერა:</b> _________________________</p>
        <script>window.onload=function(){{setTimeout(function(){{window.print()}},500)}}</script>
        </body></html>'''
        return html

    @app.route('/')
    def index():
        return render_template('form_troponin.html', template=TROPONIN_TEMPLATE)

    @app.route('/trop/print', methods=['POST'])
    @app.route('/generate_trop_pdf', methods=['POST'])
    def generate_trop_pdf():
        fd = request.form.to_dict()
        html = create_print_html(fd)
        return Response(html, mimetype='text/html')

    @app.route('/generate_trop_doc', methods=['POST'])
    def generate_trop_doc():
        fd = request.form.to_dict()
        doc = create_troponin_document(fd)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        filename = f"Troponin_{fd.get('last_name', '')}_{datetime.now().strftime('%Y%m%d')}.docx"
        return send_file(buffer, as_attachment=True, download_name=filename,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    print("❤️ Troponin სერვერი: http://127.0.0.1:5004")
    app.run(debug=False, host='127.0.0.1', port=5004, use_reloader=False)


# ====== MAIN აპლიკაცია (პორტი 8080) ======
def run_main():
    os.chdir(get_base_path())
    from flask import Flask, render_template, redirect, jsonify

    base_path = get_base_path()
    template_folder = os.path.join(base_path, 'templates')
    app = Flask(__name__, template_folder=template_folder)

    @app.route('/')
    def index():
        return render_template('index.html')

    @app.route('/cbc')
    def route_cbc():
        return redirect('http://127.0.0.1:5000/')

    @app.route('/urine')
    def route_urine():
        return redirect('http://127.0.0.1:5001/')

    @app.route('/crp')
    def route_crp():
        return redirect('http://127.0.0.1:5002/')

    @app.route('/lipid')
    def route_lipid():
        return redirect('http://127.0.0.1:5003/')

    @app.route('/trop')
    def route_trop():
        return redirect('http://127.0.0.1:5004/')

    @app.route('/signature/status')
    def sig_status():
        return jsonify({"exists": False, "url": ""})

    @app.route('/search')
    def search_patients():
        return jsonify({"results": []})

    print("🏠 მთავარი სერვერი: http://127.0.0.1:8080")
    app.run(debug=False, host='127.0.0.1', port=8080, use_reloader=False)


# ====== მთავარი გაშვება ======
if __name__ == '__main__':
    freeze_support()

    print("=" * 60)
    print("🏥 PREMIUM MEDI - ლაბორატორიული სისტემა")
    print("=" * 60)
    print()

    p1 = Process(target=run_cbc)
    p2 = Process(target=run_urine)
    p3 = Process(target=run_crp)
    p4 = Process(target=run_main)
    p5 = Process(target=run_lipid)
    p6 = Process(target=run_trop)

    p1.start()
    time.sleep(1)
    p2.start()
    time.sleep(1)
    p3.start()
    time.sleep(1)
    p5.start()
    time.sleep(1)
    p6.start()
    time.sleep(1)
    p4.start()
    time.sleep(2)

    print()
    print("=" * 60)
    print("✅ ყველა სერვისი გაშვებულია!")
    print("=" * 60)
    print()
    print("📌 ბმულები:")
    print("   🏠 მთავარი:  http://127.0.0.1:8080")
    print("   🩸 CBC:      http://127.0.0.1:5000")
    print("   🧪 Urine:    http://127.0.0.1:5001")
    print("   🧬 CRP:      http://127.0.0.1:5002")
    print("   🩸 Lipid:    http://127.0.0.1:5003")
    print("   ❤️ Troponin: http://127.0.0.1:5004")
    print()
    print("=" * 60)
    print("⚠️  დასახურად დააჭირეთ Ctrl+C")
    print("=" * 60)

    webbrowser.open('http://127.0.0.1:8080')

    try:
        p1.join()
        p2.join()
        p3.join()
        p4.join()
        p5.join()
        p6.join()
    except KeyboardInterrupt:
        print("\n🛑 სერვერები ჩერდება...")
        p1.terminate()
        p2.terminate()
        p3.terminate()
        p4.terminate()
        p5.terminate()
        p6.terminate()
        print("✅ დასრულებულია")